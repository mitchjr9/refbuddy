"""
RefBuddy — Your Minnesota HS Football Referee Assistant & Film Coach
Version 1.1 — Ref Hub Restructure + Working Presets + Stronger CTAs

Changes from v1.0.1:
  - REF HUB: three sections collapsed to two. Pre-Game Meeting now leads (most
    used, needs no upload). Crew Eval and Ref Eval were near-identical forms
    differing only in scope, so they are merged into one Evaluations section
    with a scope dropdown — "Full Crew" routes to CREW_EVAL_PROMPT, any single
    position routes to REF_EVAL_PROMPT. Both now accept photos as well as video.
  - FIX: Film & Grade quick presets did nothing. The buttons set a local var
    passed as value= to a text_area that already had key="fg_q" — Streamlit
    ignores value= on rerun when widget state exists. Presets now write to
    st.session_state["fg_q"] and rerun.
  - CTA: primary buttons (Analyze, Run RefGrade, Generate Agenda, active Ref Hub
    section) render solid #003087 with white text via a new button[kind=primary]
    rule, instead of blending in with every other light button.
  - Claude logo 70px → 85px; assignor notes placeholder simplified.

Version 1.0.1 — Light Theme Pinned (root-cause fix)

THE FIX THAT MATTERED: added .streamlit/config.toml with base = "light".
Until now the app shipped with no theme config, so Streamlit auto-detected the
visitor's OS/browser colour scheme. Anyone in dark mode got Streamlit's DARK
BaseWeb tokens applied underneath our custom CSS — the true source of the
invisible dropdown values, black expander headers, dark uploader dropzone and
dark password field. Every earlier CSS layer was treating symptoms.

Also fixed: the selectbox rule had been applying background-color to EVERY
descendant of the control, which painted over the selected-value text. Colour
now applies to the subtree; background applies only to the outer control, with
inner value containers explicitly transparent.

Version 1.0 — PUBLIC RELEASE

Shipped at refbuddy.ai for MSHSL football officials, August 2026.

Four tabs: Home chat, Film & Grade, Ref Hub, Quiz.
Backed by the RefBuddy Knowledge Base — NFHS rule facts (2021 baseline plus all
2022-2026 changes), MSHSL Minnesota modifications, crew and position mechanics,
and multiple seasons of veteran officials' game notes.

Production posture:
  - Shared-password access gate (APP_PASSWORD), fails closed on Render
  - Per-session analysis cap (400 frames) plus an account-level spend limit
  - Prompt caching on the knowledge base — repeat calls ~90% cheaper
  - Copyright guardrail: the model never reproduces NFHS/MSHSL text verbatim,
    only plain-language summaries with rule numbers
  - Terms of Use, non-affiliation notice, and no-warranty disclaimer in footer

Final polish in this release:
  - "Assignor Hub" renamed "Ref Hub"; Individual Referee Evaluation now uses a
    football icon
  - Pre-Game Meeting description rewritten to reference the RefBuddy Knowledge
    Base rather than an internal variable name
  - NEW CSS Layer 4: fixes BaseWeb form controls that ignored the outer theme —
    selectbox values rendering white-on-white, multiselect chips, password
    fields rendering dark-on-dark, and the truncated "Press Enter to apply"
    overlay. Uses -webkit-text-fill-color, which overrides `color` on inputs in
    WebKit browsers and is required for the fix to hold in Safari and Chrome.
  - Claude logo enlarged to 60px

Prior — Version 3.6 — UI Polish (logo size, expander/uploader contrast, Quiz rename)

Changes from v3.5:
  - Claude logo enlarged 26px → 46px; "Powered by" bumped to 1.05rem and
    vertically centred against it.
  - Removed the "Prompt caching active" sidebar caption (caching still on —
    it just does not need to be user-facing).
  - NEW CSS Layer 3: pins expanders and the file uploader to the light theme.
    Root cause of the black bars was that the old ".streamlit-expanderHeader"
    class no longer exists in current Streamlit; those components fell through
    to Streamlit's own dark tokens on hover/focus/open. Now targeted via
    [data-testid="stExpander"] and [data-testid="stFileUploaderDropzone"]
    across every interaction state, including the chevron icon and the
    "200MB per file" helper text.
  - "Quiz & Drills" renamed to "Quiz" (tab label, heading, home chip) and the
    description reworded to drop the internal CORE_KNOWLEDGE reference.

Prior — Version 3.5 — Film & Grade Merge + Photo Upload + Claude Branding

Changes from v3.4:
  - TABS 5 → 4: Game Film and RefGrade merged into a single "🎬 Film & Grade"
    tab. One upload feeds two modes (Ask a Question / RefGrade Evaluation),
    removing the confusion of two similar-looking film tabs.
  - PHOTO UPLOAD: the uploader now accepts jpg/jpeg/png alongside mp4/mov.
    image_to_frame_b64() converts stills into the same base64-JPEG format
    extract_frames() produces, so photos and video frames share one pipeline.
    Photos are now the primary path — video uploads are unreliable on phone
    data and large clips often fail.
  - BRANDING: sidebar now shows the "Powered by" + Claude logo lockup
    (Claude.png, loaded via _asset_data_uri with a text fallback if missing),
    replacing the model pill and "Powered by Anthropic" caption.
  - COPY: "Built by a ref, for refs" (comma added, all four locations).
    Knowledge Base line and footer reworded to "Years of NFHS veteran
    officials' game notes and NFHS/MSHSL rulebook facts and interpretations"
    — de-risks the earlier phrasing that implied reproducing the rulebooks.
    Footer disclaimer now says "Not official NFHS/MSHSL interpretation."

Prior — Version 3.4 — Secrets Resolution Fix (env vars) + Fail-Closed Auth

Changes from v3.3:
  - FIX: StreamlitSecretNotFoundError crash on Render. st.secrets does NOT read
    environment variables, and touching it with no secrets.toml on disk raises
    StreamlitSecretNotFoundError — not a KeyError — so the old narrow except
    clause let it through and killed the app at boot.
  - NEW get_secret(): resolves os.environ first, then secrets.toml, then a
    legacy [anthropic] table. Every st.secrets access wrapped in except Exception.
  - NEW running_on_render(): detects Render via its RENDER=true variable.
  - FAIL-CLOSED AUTH: missing APP_PASSWORD on a Render host now stops the app
    with an explanatory error instead of silently serving it unprotected.
    Locally, a missing password still just skips the gate.

Prior — Version 3.3 — Prompt Caching + GSOA 6/17/26 Notes

Changes from v3.2:
  - PROMPT CACHING: CORE_KNOWLEDGE hoisted out of all six prompt strings into
    a single CACHED_KB_BLOCK with cache_control ephemeral, placed FIRST in the
    system parameter. All prompts (chat, RefGrade, quiz, crew eval, ref eval,
    pre-game) now share ONE cache entry. system_blocks() helper wires it into
    all five API call sites. Cache reads bill at 0.10x base input rate and do
    not count toward the input-tokens-per-minute rate limit.
  - CORE_KNOWLEDGE Section 5B: GSOA meeting notes 6/17/26 — inadvertent
    whistles, kneel-down mechanics, helmet-off rule and its foul exception,
    illegal shift as live-ball foul, roughing passer vs late hit OOB, blocking
    below waist, tack-on enforcement.
  - CORE_KNOWLEDGE Section 5C: 2025 midseason review — pre-game checklist,
    backward pass signal, weekly study rotation.

Prior — Version 3.2 — Public Launch Hardening (Password Gate + Usage Cap + Legal)

Changes from v3.1:
  - ACCESS GATE: shared-password login via st.secrets["APP_PASSWORD"], using
    hmac.compare_digest. Skipped if APP_PASSWORD unset (local dev convenience).
  - SESSION USAGE CAP: MAX_FRAMES_PER_SESSION = 400. spend_frames() charges the
    budget before every vision call (Film, RefGrade, Crew Eval, Ref Eval).
    Sidebar shows remaining frames.
  - COPYRIGHT GUARDRAIL in SYSTEM_PROMPT: never reproduce verbatim NFHS/MSHSL
    text; decline "quote the rule" requests; answer with plain-language summary
    plus rule number only.
  - TERMS OF USE / DISCLAIMER expander in footer.
  - MSHSL 2026 season calendar added to CORE_KNOWLEDGE Section 2F.

Prior — Version 3.1 — Hybrid CORE_KNOWLEDGE: 2021 Rulebook + 2022-2026 Changes

Changes from v3.0:
  - CORE_KNOWLEDGE rebuilt with new Section 0 "2022-2026 NFHS Rules Changes"
    at the top, containing all official changes from every project document:
    2022-2024_rulechanges_POE_Football.md, 2024nfhsfootballruleschangefinaljune2024.pdf,
    2025 FB Rules Changesfinalapril2025.pdf, 2025 NFHS POE (FINAL 2/18/25),
    2026_NFHS_Football_Rules_and_Editorial_Changes (2/18/26)
  - Each change includes: rule number, year, old language, new language, why it matters
  - Quick-reference table listing every changed rule from 2021 baseline
  - SYSTEM_PROMPT updated with CRITICAL LAYERING RULE: default to 2021 rulebook
    unless Section 0 overrides it; cite the year when a change applies
  - Intentional grounding 2022/2023 two-step change fully documented
  - Defenseless player 2023 expansion documented
  - Forward fumble OOB 2025 rule fully documented with beanbag note
  - All v3.0 features (secrets-only API key, hard-coded Sonnet, light dropdowns) preserved

Tabs: 🏈 Home | 🎬 Film & Grade | 👥 Assignor Hub | 📝 Quiz
Run:  streamlit run app_v2.py
"""

# ── Standard library ──────────────────────────────────────────────────────────
import base64
import datetime
import hmac
import json
import os
import subprocess
import sys
import tempfile
import urllib.parse

# ── Third-party ───────────────────────────────────────────────────────────────
import anthropic
import numpy as np
import streamlit as st

# ── OpenCV — auto-install if missing ─────────────────────────────────────────
try:
    import cv2
    OPENCV_AVAILABLE = True
except ImportError:
    try:
        subprocess.check_call(
            [sys.executable, "-m", "pip", "install", "opencv-python-headless", "-q"]
        )
        import cv2
        OPENCV_AVAILABLE = True
    except Exception:
        OPENCV_AVAILABLE = False


# =============================================================================
# CORE KNOWLEDGE BASE
# =============================================================================

CORE_KNOWLEDGE = """
# RefBuddy Core Knowledge Base
## Minnesota High School Football — Referee Reference

---

## 0. 2022–2026 NFHS RULES CHANGES & UPDATES

> **INSTRUCTION FOR ALL RESPONSES:** Default to the 2021 NFHS Rulebook (Sections 1–8 below) for any question unless a specific update in THIS section overrides it. Always cite the year when a change applies. If a rule changed in 2022 and then was further clarified in 2023, note both years.

---

### 2022 Rule Changes

**[2022] Jersey Number "0" Legalized — Rule 1-4-3**
- **Old:** Numbers 1–99 only.
- **New:** Players may now wear the number 0.
- **Why it matters:** Squad number assignments and ineligible receiver detection — #0 is treated as eligible (1–49 or 80–99 range does not include 0, so #0 IS eligible by position unless declared ineligible). Confirm with your state interpreter; MSHSL has adopted this.

**[2022] Chop Block Redefined — Rule 2-3-8**
- **Old:** A chop block was defined as a block at or below the knees of an opponent being blocked above the waist by a teammate.
- **New:** A chop block is a combination block by two or more offensive teammates against the same opponent (not the runner) where at least one block is above the waist and at least one is below the waist.
- **Why it matters:** Broader definition — any high-low combination on the same non-runner defender is now a chop block. 15-yard penalty. Look for OL double-teaming with cut blocks on defensive linemen.

**[2022] Intentional Grounding — Outside-Pocket Exception Added — Rule 7-5-2**
- **Old (2021):** Intentional grounding required an eligible receiver in the area of the pass. No pocket exception existed in NFHS (unlike NCAA).
- **New (2022):** A passer who is outside the tackle box (outside the normal tackle position on either side) may legally throw an incomplete forward pass that reaches the line of scrimmage without an eligible receiver in the area — NO foul.
- **Why it matters:** MAJOR change aligning NFHS closer to NCAA on scrambling QBs. If the QB rolls out past the tackle position AND the ball reaches the neutral zone, it is NOT intentional grounding even with no receiver in the area. The 2021 rulebook has NO pocket exception — apply 2022 rule instead.
- **Penalty if foul still committed:** 5 yards from spot of pass + loss of down.

**[2022] Ball Change Authority Expanded — Rule 1-3-3**
- **Old:** Only the referee could order a ball change between downs.
- **New:** Any game official may order a ball change.
- **Why it matters:** Crew efficiency — wing officials and BJ can now act on a damaged or unfit ball without flagging the referee.

**[2022] Game Clock Start Option After Fouls (Last 2 Minutes) — Rule 3-4-7**
- **Old:** No specific option for offended team regarding clock start after accepting a penalty in the final 2 minutes.
- **New:** In the last two minutes of either half, when a foul is accepted, the offended team may choose to have the game clock start on the snap (rather than on the ready signal).
- **Why it matters:** Significant clock-management tool for trailing teams. If defense commits a foul with <2:00 left, offense can choose snap-start, which prevents the defense from using the clock stoppage to recover time.

**2022 Points of Emphasis:** Targeting and defenseless player protection; legal uniforms and equipment verification; sportsmanship.

---

### 2023 Rule Changes

**[2023] Intentional Grounding — Snap Receiver Restriction — Rule 7-5-2**
- **Old (2022):** Any passer outside the pocket could use the neutral-zone exception.
- **New (2023):** ONLY the player who directly receives the snap (usually the QB) may legally use the outside-pocket exception to throw the ball away. A player who received a handoff or lateral CANNOT use this exception.
- **Why it matters:** Closes a potential loophole where a receiver who took a lateral behind the LOS could spike the ball. If a running back or wide receiver receives a pitch/lateral and then throws the ball away out of bounds short of the LOS, it IS intentional grounding — 5 yards + loss of down.

**[2023] Defenseless Player Definition Expanded — Rule 2-32-16**
- **Old (2021):** Defenseless players included passer, receiver attempting catch, returner, runner whose progress is stopped, player on ground, player out of play.
- **New (2023):** A defender who is in the act of or just completed an interception attempt is now explicitly classified as defenseless. Added language: a player attempting to intercept a pass is as vulnerable as a receiver catching one.
- **Why it matters:** Hits on defensive backs going for the ball must be evaluated under defenseless-player rules — open-hands contact or wrap-up required. Non-open-hands forceful contact = 15-yard personal foul.

**[2023] Inbounds Re-establishment — Rule 2-29-1**
- **Old:** General requirement to re-establish inbounds — specific language was ambiguous.
- **New:** A player who goes out of bounds must re-establish both feet (or one knee/other body part) inbounds before legally touching a forward pass or being the first to touch a kick. One foot is not sufficient after going OOB.
- **Why it matters:** Receiver who runs out on a route and comes back — if only one foot re-established, any catch is incomplete. Both feet must touch inbounds before the catch.

**[2023] Penalty Enforcement — Previous Spot Clarified — Rules 10-4 and 10-6**
- **Old:** Enforcement spots for fouls behind the LOS were sometimes inconsistently applied.
- **New:** Clarified that most fouls occurring behind the line of scrimmage during running plays or loose-ball plays are enforced from the previous spot (not from the spot of the foul).
- **Why it matters:** Prevents confusion on off-tackle runs where a blocker holds 5 yards behind the LOS — the 10-yard penalty comes from the previous spot, not the spot of the hold, which could otherwise result in a net gain for the offense.

**[2023] Towel Rules Clarified — Rule 1-5-3a**
- **Old:** Towel rules were less specific regarding color and logo restrictions.
- **New:** Solid-colored towels are allowed (except brown or penalty-flag yellow/gold colors). Logos and manufacturer's marks are allowed but limited in size. Towels must be tucked into waistband; no flags hanging.
- **Why it matters:** Equipment check item — any brown or flag-colored towel must be removed before the game. Logo towels are legal as long as they meet size limits.

**[2023] Intentional Pass Interference — No Longer Automatic USC — Rule 7-5 Penalty**
- **Old:** Intentional pass interference (a defender deliberately fouling to prevent a TD) carried a 15-yard PI penalty PLUS an automatic unsportsmanlike conduct foul (15 yards + possible DQ).
- **New:** The automatic USC has been removed. Intentional PI is still 15 yards from the previous spot. Officials may still flag USC separately if the conduct warrants it, but it is no longer automatic.
- **Why it matters:** One less automatic double-foul situation. The foul is still serious — 15-yard PI from previous spot is significant — but DQ is no longer a default outcome.

**2023 Points of Emphasis:** Helping the runner; coach-official communication; game management.

---

### 2024 Rule Changes

**[2024] Home Team Uniform — Same Dark Color Required — Rule 1-5-1b(3)**
- **Old:** Home team wore dark jerseys, but no specific requirement that all jerseys be the same shade.
- **New:** The jerseys of the home team shall ALL be the same dark color(s) that clearly contrasts with white. Mixing different dark shades (e.g., half the team in navy, half in royal blue) is a violation.
- **Why it matters:** Pre-game equipment check — if the home team's jerseys are not all the same color family, notify the head coach before kickoff. This is an equipment violation (Rule 1-5-1), not a uniform number issue.

**[2024] Number Body Color — Single Solid Contrasting Color — Rule 1-5-1b(3)**
- **Old (through 2023):** Number body could be either a contrasting color OR the same color as the jersey with a contrasting border (two options).
- **New (effective 2024):** The entire body of the number (horizontal bars and vertical strokes) shall be a SINGLE SOLID COLOR that clearly contrasts with the jersey body color. The dual-color option is gone.
- **Why it matters:** Jersey number legibility check at equipment inspection.

**2024 Points of Emphasis:** Sportsmanship and altercation prevention; player equipment and enforcement; formations (illegal formations and knee pad compliance).

---

### 2025 Rule Changes

**[2025] Forward Fumble Out of Bounds — Spot of Fumble Rule — Rules 3-4-2d (NEW), 4-3-1 EXCEPTION (NEW), 8-5-2a EXCEPTION**
- **Old (2021):** When a fumble went out of bounds, the ball was returned to the spot where it crossed the sideline/end line.
- **New (2025):** When a FORWARD fumble goes out of bounds (or is ruled out of bounds between the goal lines), the ball is returned to the SPOT OF THE FUMBLE — not where it went out of bounds.
- **Why it matters:** HUGE game-management change. A QB who fumbles at the LOS and the ball rolls 15 yards downfield out of bounds — the offense gets the ball back at the LOS where the fumble occurred, NOT at the 15-yard spot. Use a beanbag to mark the spot of the fumble. Note: this applies to FORWARD fumbles only; backward fumbles still go where they went OOB.
- **Personal note:** This is a 2025 new rule — the 2021 rulebook language does NOT apply to forward fumbles OOB anymore.

**[2025] Electronic Signs Allowed — Rule 1-5-3c(2)**
- **Old:** No electronic devices allowed to relay plays from sideline to players.
- **New:** Fixed electronic signs with play signals (non-audio, non-video) are allowed for relaying plays from the sideline. In-helmet communication is still prohibited. Players may not watch video between plays.
- **Why it matters:** Many teams now use electronic signaling systems. Legal as long as: no audio, no video feed to players, and no in-helmet communication device.

**[2025] No Audio/Video Recording Devices on Players — Rule 1-5-3c(3) NEW**
- **Old:** No specific rule against players wearing recording/transmission devices.
- **New:** No player participating in the game may wear any audio or video device to record or transmit audio or video.
- **Why it matters:** New equipment violation — any body camera, GoPro, or similar device worn by a player on the field is illegal. Must be removed before participation. Enforce as illegal equipment.

**[2025] Tooth/Mouth Protector Attachment Restrictions — Rule 1-5-1d(5)a (Effective 2026 season)**
- **New:** Items attached to the tooth and mouth protector that do not serve a protective function or that pose a health/risk issue are not allowed.
- **Why it matters:** Flavored candy attachments, decorative items on mouthguards = illegal starting 2026.

**2025 Points of Emphasis (Official — Final 2/18/25):**
1. **Illegal and Improperly Worn Player Equipment** — Illegal = prohibited items (jewelry, tinted visor, non-compliant eyeshade, bands on arm/neck/legs, back pads uncovered, non-conforming towels). Improperly worn = legal equipment not worn as designed (pants not covering knees, mouthguard not in at snap start, shoulder pads not covered by jersey). Player removed for one down for improperly worn equipment. Chronic violations = delay-of-game or USC on head coach.
2. **Sportsmanship** — Penalize WITHOUT WARNING: gun gestures, throat slashes, rehearsed poses, sexual gestures, dunking ball over crossbar, removing helmet to celebrate/protest, dancing, somersaults/flips, spiking or spinning the ball. Any act intended to taunt, demean, or disrespect an opponent or the game.
3. **Defenseless Player / Targeting** — A player is defenseless based on their OWN physical position and focus, not another player's action. Contact on a defenseless receiver is limited to: incidental contact while playing the ball, contact initiated with open hands, or a wrap-up tackle attempt. All other forceful contact = 15-yard personal foul. Targeting: takes aim and initiates contact above the shoulders. NOT automatic DQ in NFHS (unlike NCAA). Game officials: when in doubt, throw the flag — be supported in doing so.

---

### 2026 Rule Changes

**[2026] Play Cards on Forearm/Wrist AND Belt — Rule 1-5-3c(9)**
- **Old:** Play cards allowed on forearm/wrist area only.
- **New:** Play cards may now also be worn on the belt area of the body.
- **Why it matters:** Equipment check item — both locations now legal. No penalty for cards worn on the belt.

**[2026] Slap to Head — NEW Personal Foul — Rule 9-4-7 (NEW)**
- **Old:** No specific rule against using hands to slap an opponent's head (covered under general illegal use of hands provisions).
- **New:** No player may use hand(s) to slap an opponent's head. This is a new, specific 15-yard personal foul.
- **Why it matters:** Now a distinct and named foul — not just "illegal use of hands." If you see a player slap an opponent's helmet, it's Rule 9-4-7, 15 yards. Flag it immediately. Does not require twist or pull like facemask — any slap to the head is the foul.

**2026 Points of Emphasis:**
1. **Flagrant and Unsportsmanlike Fouls** — Consistent enforcement. Flagrant fouls = DQ. No tolerance for taunting or postgame confrontations.
2. **Helping the Runner** — Pushing/pulling the runner by a teammate is illegal. Watch for OL pulling QB over the pile on short yardage.
3. **Sideline Management and Control** — Coaches' restricted area enforcement. Team box limits (MSHSL: 15-yard lines). Non-players must stay in the team box.
4. **Identification of the NFHS Authenticating Mark on Game Balls** — Verify the authenticating mark on all game balls before the game, especially at 9th grade and above (Rule 1-3-1). Balls without the mark are not legal for varsity competition.

---

### Quick-Reference: Rules That CHANGED from 2021 Baseline

| Rule | Year | What Changed |
|------|------|-------------|
| 1-4-3 | 2022 | Jersey #0 now legal |
| 2-3-8 | 2022 | Chop block = any high-low combination on same non-runner |
| 7-5-2 | 2022 | Outside-pocket exception added to intentional grounding |
| 7-5-2 | 2023 | Outside-pocket exception limited to snap receiver only |
| 2-32-16 | 2023 | Defender attempting interception now explicitly defenseless |
| 2-29-1 | 2023 | Both feet must re-establish inbounds (not one foot) |
| 1-5-3a | 2023 | Towel color/logo rules clarified |
| 7-5 PENALTY | 2023 | Intentional PI no longer carries automatic USC |
| 1-5-1b(3) | 2024 | Home team jerseys all must be same dark color |
| 1-5-1b(3) | 2024 | Number body must be single solid contrasting color |
| 3-4-2d | 2025 | Forward fumble OOB returns to spot of fumble (not OOB spot) |
| 1-5-3c(2) | 2025 | Electronic play signs (non-audio/video) allowed |
| 1-5-3c(3) | 2025 | No audio/video recording devices on players |
| 1-5-3c(9) | 2026 | Play cards now allowed on belt area in addition to wrist |
| 9-4-7 | 2026 | Slap to opponent's head = new 15-yard personal foul |

---

---

## 1. NFHS RULE HIERARCHY & KEY CITATIONS

### Rule 1 — The Game, Field, Players and Equipment
- **Game clock**: 48 minutes, four 12-minute periods (Rule 3-1-1)
- **Ball specs**: 14-15 oz, 12.5-13.5 psi, NFHS Authenticating Mark required 9th grade+ (Rule 1-3-1)
- **Mandatory equipment**: Helmet (NOCSAE), face mask, jersey, hip/knee/shoulder/thigh pads, mouthguard, pants covering knees, shoes (Rule 1-5-1)
- **Illegal equipment always includes**: jewelry, tinted eyeshields, ball-colored items, non-NOCSAE gloves (Rule 1-5-3)
- **2025 POE**: Illegal & improperly worn equipment top emphasis. Pants not covering knees, mouthguard not in at snap, shoulder pads not covered = improperly worn. Player sent off for one down (Rule 1-5-5). Chronic violations = delay of game or USC on head coach.
- **2025 NEW**: Electronic signs allowed to relay plays (non-audio, fixed) [Rule 1-5-3c(2)]. No audio/video devices on players [1-5-3c(3) NEW].
- **2026 NEW**: Play cards now allowed on forearm/wrist AND belt area [1-5-3c(9)].
- **2026 NEW**: No player may use hand(s) to slap opponent's head — 15-yard penalty [Rule 9-4-7].

### Rule 2 — Definitions
- **Free-Blocking Zone (FBZ)**: Rectangular area 4 yards either side of snap spot, 3 yards behind each LOS. Disintegrates when ball leaves zone (Rule 2-17-1,4)
- **Blocking below waist in FBZ**: Legal only if all players in zone at snap, contact in zone, immediate action (Rule 2-17-2)
- **Block in back in FBZ**: Legal only by offensive linemen on LOS, against defensive players in zone, contact in zone (Rule 2-17-3)
- **Blindside block**: Block against opponent who does not see blocker approaching. Outside FBZ with forceful contact must be initiated with open hands or it's a foul (Rule 2-3-10, 9-4-3n)
- **Targeting**: Any player takes aim and initiates contact above shoulders with helmet, forearm, hand, fist, elbow, or shoulders (Rule 2-20-2). NOT automatic DQ in NFHS (unlike NCAA).
- **Spearing**: Crown of helmet to opponent at shoulders or below (Rule 2-20-1c)
- **Defenseless player** (Rule 2-32-16): passer, receiver attempting catch, intended receiver during/after interception, runner whose progress stopped in grasp, kick returner attempting to catch/recover, player on ground (including slider), player out of play, player receiving blindside block without open hands. When in question, player IS defenseless. 2025 POE top emphasis.
- **Forward fumble out of bounds** (2025 NEW Rule 3-4-2d): Ball returned to SPOT OF FUMBLE, not where it went out of bounds.
- **Scrimmage kick formation**: No player in hand-to-hand snap position; either (a) knee-down holder 7+ yards back with kicker 3 yards behind, or (b) player 10+ yards back to receive snap (Rule 2-14-2)
- **Fair catch signal**: Extending and laterally waving ONE arm at full arm's length above head (Rule 2-9-3)

### Rule 3 — Periods, Time Factors and Substitutions
- **Play clock 25 seconds**: starts on ready-for-play signal — prior to try after score; start of period/OT; after inadvertent whistle; after charged timeout; after official's timeout; after legal kick when team awarded new series; after referee stops clock (Rule 3-6-1a(1))
- **Play clock 40 seconds**: starts when ball declared dead after a down — all other situations (Rule 3-6-1a(2))
- **Illegal substitution**: 5-yard penalty if 12 in huddle. KILL IT IMMEDIATELY — if not killed, becomes 15-yard illegal participation (Rule 3-7 PENALTY, 9-6-4c). Personal note: KILL 12 PLAYERS ON DEF RIGHT AWAY.
- **Halftime**: 15 minutes normal; min 10, max 20 with notice (Rule 3-1-6). Mandatory 3-minute warm-up after intermission.
- **Overtime (NFHS)**: Ball at opponent's 10-yard line, 1st and goal, each team gets series. Defense gains possession = ball dead immediately.

### Rule 4 — Ball in Play, Dead Ball, Out of Bounds
- **Ball dead when** (Rule 4-2-2): runner out of bounds / forward progress stopped / any body part except hand or foot down; loose ball out of bounds; forward pass incomplete; kick breaks plane of R's goal line; helmet completely off runner; score; inadvertent whistle
- **Inadvertent whistle** (Rule 4-2-3): Ends the down. If during snap/pass in flight or legal kick = down replayed. If during loose ball = last team in possession chooses. If during player possession = team in possession chooses dead-ball spot or replay.
- **2025 CHANGE**: Forward fumble OOB = returned to SPOT OF FUMBLE (not where it went out).

### Rule 5 — Series of Downs
- **New series awarded to R** if K touches scrimmage kick beyond expanded neutral zone before R (first touching). R can take ball at spot of first touching or result of play (Rule 5-3-3g)
- **Line to gain**: 10 yards from foremost point of ball when series begins. If extends into end zone, goal line is line to gain (Rule 5-3-1,2)

### Rule 6 — Kicking
- **Kickoff**: From K's 40-yard line; at least 4 K players on each side of kicker; no K player more than 5 yards behind free-kick line except kicker (Rule 6-1-1,3,4)
- **Free kick after safety**: From K's 20-yard line; may punt, place kick, or drop kick (Rule 6-1-1b, 6-1-2)
- **Pop-up kick**: ILLEGAL (Rule 6-1-11)
- **Free kick out of bounds** R options (Rule 6-1-9): (a) 5-yard penalty + re-kick; (b) 5-yard penalty from succeeding spot; (c) ball 25 yards beyond previous spot; (d) decline + inbounds spot
- **Touchback**: Ball snapped at R's 20-yard line NFHS (NCAA = 25-yard line) (Rule 8-5-4)
- **Fair catch** (Rule 6-5): After fair catch, offense may free kick OR snap at spot (NFHS only). K cannot contact R or obstruct path to ball.
- **First touching**: K touches scrimmage kick beyond expanded neutral zone before R. Mark with beanbag. Don't wind clock.
- **K recovering own kick**: Ball dead at spot. K cannot advance recovered kick or muffed punt.
- **Onside kick**: Kick must travel at least 10 yards before K can recover unless R touches first.

### Rule 7 — Snapping, Handling, Passing
- **Legal formation** (Rule 7-2-5): At least 5 players numbered 50-79 on LOS; no more than 4 backs; only 1 player may penetrate vertical plane of nearest lineman's waist
- **Only 1 player in motion** at snap, only if NOT toward opponent's goal line (Rule 7-2-7)
- **Illegal shift**: All A players must stop simultaneously for at least 1 second before snap (Rule 7-2-6)
- **False start** (Rule 7-1-7): Feigned charge simulates snap; any act to cause B to encroach; A lineman between snapper and end of line moves hand or makes quick movement
- **Forward pass legal**: Both passer's feet in or behind neutral zone when released; only one per down; A player only (Rule 7-5-1)
- **Intentional grounding**: 5-yard penalty + loss of down from spot of pass. NFHS: illegal if no eligible receiver in area REGARDLESS of pocket (Rule 7-5-2)
- **Ineligible downfield**: OL cannot advance beyond expanded neutral zone before legal forward pass crossing neutral zone. 5-yard penalty (Rule 7-5-12). "A good 3 yards" per personal notes.
- **Eligible receivers**: Players on ends of LOS or legally behind LOS, numbered 1-49 or 80-99. All B players eligible (Rule 7-5-6)
- **Pass interference**: Only beyond neutral zone; must have physical contact that IMPEDES the other player's opportunity. 15-yard penalty (Rule 7-5-7,10)

### Rule 8 — Scoring
- **Scoring values**: TD=6; PAT kick=1; PAT run/pass=2; FG=3; Safety=2 (Rule 8-1)
- **Touchdown**: Ball penetrates the vertical plane of the opponent's goal line (Rule 8-2-1a)
- **Try**: Snapped from B's 3-yard line (Rule 8-3-1). Kick=1 pt; run or pass=2 pts. Only A can score. FG attempt behind LOS = live ball until B secures possession.
- **Field goal**: Must pass between uprights (inside) and above crossbar. Blocked FG in bounds = LIVE BALL. Unsuccessful FG crossing goal line = touchback (NFHS).
- **Safety**: 2 points; followed by free kick from K's 20-yard line (Rule 8-5-2, 6-1-1b)
- **Personal foul on TD**: TD counts. Offense = enforce on ensuing kickoff. Defense = enforce on try OR ensuing kickoff (Rule 8-2-2,3)

### Rule 9 — Conduct of Players and Others
- **Holding** (Rule 9-2-1c, 9-2-3c): Hook, lock, clamp, grasp, encircle, or hold to restrain. 10-yard penalty.
- **Illegal use of hands** (Rule 9-2-1a,3a): Technique not permissible; swinging/throwing elbow or forearm faster than shoulders; initiating contact above opponent's shoulders. 10-yard penalty.
- **Hands to face**: Placement on helmet = 5 yards; violent punch/grab = 15 yards.
- **Blocking below waist** (Rule 9-3-2): Illegal except in FBZ or to tackle runner. 15-yard penalty.
- **Block in back** (Rule 9-3-5): Illegal except in FBZ or using hands/arms above waist to ward off. 10-yard penalty.
- **Chop block**: Combination block where one is low (at or below knee) and one is high. ALWAYS illegal. 15 yards (Rule 2-3-8, 9-3-6).
- **Blindside block outside FBZ with forceful contact not initiated with open hands**: 15-yard penalty (Rule 9-4-3n)
- **Facemask**: Incidental grasp without twist/turn/pull = 5 yards. Grasp + twist/turn/pull = 15 yards (Rule 9-4-3h)
- **Horsecollar**: Grab inside back/side collar or name plate area and pull to ground = 15-yard penalty. From end of run if past LOS; down NOT replayed (Rule 9-4-3k). Personal note: personal game MSHSL Mound Westonka v Annandale.
- **Targeting**: Takes aim, initiates contact above shoulders. 15-yard penalty. NOT automatic DQ in NFHS (Rule 9-4-3m, 2-20-2). 2025 POE major enforcement emphasis.
- **Spearing**: Crown of helmet to opponent at shoulders or below. 15-yard penalty + possible DQ if flagrant (Rule 9-4-3i, 2-20-1c)
- **Defenseless player contact** (2025 POE): Contact must be incidental, open hands, OR wrap-up tackle attempt. Non-open-hands forceful contact = foul. 15-yard penalty.
- **Roughing the passer**: 15 yards + automatic first down (Rule 9-4-4)
- **Roughing the kicker/holder**: 15 yards + automatic first down from previous spot (Rule 9-4-5)
- **Running into kicker**: Displaces without roughing. 5-yard penalty, previous spot.
- **Roughing the snapper**: 15 yards + automatic first down (Rule 9-4-6)
- **2026 NEW - Slap to head**: No player may use hand(s) to slap opponent's head. 15-yard penalty (Rule 9-4-7 NEW).
- **Unsportsmanlike conduct**: 15 yards. Second USC = disqualification (Rule 9-5-1). 2025 POE: penalize WITHOUT WARNING for gun gestures, throat slashes, rehearsed poses, removing helmet to celebrate, dancing, somersaults, spiking/spinning ball.
- **Illegal participation - 12 on field** (Rule 9-6-4c): 15-yard penalty NFHS (NCAA = 5-yard illegal substitution). Personal note: KILL 12 PLAYERS ON DEF RIGHT AWAY.

### Rule 10 — Enforcement of Penalties
- **Basic spot = previous spot**: fouls simultaneous with snap/free kick; fouls during loose-ball play (Rule 10-4-2)
- **Basic spot = end of related run**: fouls during running plays (Rule 10-4-4)
- **All-but-one principle**: All fouls from basic spot EXCEPT offensive foul behind basic spot during loose-ball or running play = penalized from spot of foul (Rule 10-6)
- **Dead-ball fouls**: Always from succeeding spot. Cannot combine with live-ball fouls for double foul (Rule 10-4-5)
- **Double foul**: Both teams foul during same live-ball period. Penalties cancel, down replayed (Rule 10-2-1)
- **Half the distance**: Penalty cannot take ball more than half the distance to offending team's goal line (Rule 10-1-5)
- **Roughing the passer special enforcement**: 15 yards from dead-ball spot when beyond neutral zone with no change of possession. Always automatic first down.
- **Automatic first down fouls** (Rule 10-1-7): Roughing kicker/holder; roughing passer; roughing snapper.
- **Offensive penalty behind LOS**: ENFORCED AT LOS, not spot of foul.
- **Under 2 minutes**: Offended team may choose game clock to start on snap when penalty accepted (Rule 3-4-7)

---

## 2. MSHSL MINNESOTA-SPECIFIC MODIFICATIONS (Updated 9/9/2025)

### A. Play Clock
- Visible 25/40-second play clock MAY be used at home team's discretion. Must have operator, available to both teams.
- When visible play clock present, Back Judge signals (10 sec = hand up; 5 sec = basketball chop) NOT required.

### B. Nine-Player Football
- All 15-yard penalties become 10-yard penalties if played on 80-yard field; 15 yards on 100-yard field.
- Field width: 40 yards. Hash marks: 48 feet, 4 inches from sideline.

### C. Mercy Rule (Minnesota ONLY)
- Point differential reaches 35+ points in fourth quarter = clock goes to running time.
- Clock stops ONLY for TIPS: T=Team timeouts, I=Injuries, P=Penalties, S=Scores.
- Regular timing resumes if differential drops below 30 points.
- Under mercy rule: Play ends OOB, incomplete pass, touchback, first down, change of possession, measurement = NO stop.
- Injury, Penalty, Score, Team TO = ALWAYS stop. Restarts on RFP for Injury/Penalty under mercy rule.

### D. Overtime
- Both 9-player and 11-player: "10-yard" overtime procedure from NFHS rulebook (from B's 10-yard line).
- Each team gets a possession from opponent's 10-yard line; if tied after one round, repeat.

### E. Team Boxes
- Team boxes between 15-yard lines (not 25-yard lines). Coaches' restricted area = 2-yard belt from team box to sideline.

### F. 2026 Season Calendar (MSHSL, mshsl.org)
- Season begins (first practice): **August 17, 2026**. Week 1 may only be 5 days of practice.
- First available contest date: **September 3, 2026**
- Last date for section contests: **November 6, 2026**
- End of season / Prep Bowl: **November 28, 2026**
- Maximum number of contests: 9. Official ball: Spalding Alpha.
- Teams must follow the Pre-Season Practice / Heat Acclimatization Policy.
- NOTE: As of mid-August 2026, MSHSL had not yet posted a 2026-specific Minnesota Rules
  Modifications document; the 2025 modifications (updated 9/9/2025, Section 2 above) remain
  the most recent published version. If asked about 2026 MN modifications, say the 2025
  document is the current published version and direct the user to mshsl.org to check.

### G. Weather Policy
- Lightning: suspend play minimum 30 minutes after last flash/thunder. "If you can hear it, clear it."
- Before contest: host school determines go/no-go. Once started: officials have authority (cannot be overruled by coaches, CAN be overruled by school superintendent).
- Cold: Cancel/postpone at -4F or -40F wind chill.
- Heat (WBGT): At 82.1-85F south: max 2-hour practice, restricted to helmet/shoulder pads/shorts. At 85.1-87.1F south: max 1-hour practice, no protective equipment.

---

## 3. POSITION-SPECIFIC MECHANICS

### Back Judge (BJ)
**Kickoff**: Handle ball, give to kicker, position 2 yards outside sideline on K's free-kick line (chains side). Move onto field between 9-yard mark and hash after kick. Goal line responsibility on long return. Verify K has 11 players; encroachment = kill play. Anticipate onside kick on EVERY kickoff.

**Punts**: Position between deepest receiver and chains-side sideline, ~5 yards wider, 3 yards behind. Beanbag the yard line where R takes possession. Have beanbag in hand on ALL scrimmage kicks; mark "end of kick" — NEVER beanbag forward progress. K cannot advance muffed punt. Kick breaking plane of goal line = automatic touchback. Talk to PR by name before each punt about fair catch signal.

**Scrimmage plays**: 17-20 yards beyond LOS; on end line if ball snapped inside 15-yard line. Key on strong side. Back pedal as play progresses. Play clock: raise arm at 10 seconds; basketball chop at 5 seconds (not needed when visible play clock exists). Count Team B players; straight arm and fist for 11.

**Inside 10-yard line**: Ball snapped 10-7 yards out: move to 5-yard line. Ball snapped 7 yards and in: go STRAIGHT to goal line at snap. On PAT/FG: under upright opposite from umpire. Signal YES YES YES or NO NO NO with umpire.

### Umpire (U)
**Positioning**: 12 yards behind LOS, right foot lined up on left tackle's left foot, LEFT side of QB. Watch center, left guard, left tackle. If QB rolls your way past you = QB is now your responsibility (roughing the passer). Mirror positioning with Referee on punts; 3-4 yards behind punter at edge of free blocking zone. Under 2 min: Don't spot ball until all officials set.

**Kickoff**: Line up on receiving team's line (50-yard line), 2 yards outside sideline opposite from BJ. Hold at 50 until players cross your face.

**PAT/FG**: Under upright opposite from BJ. Signal YES YES YES or NO NO NO with BJ.

### Wings (Line Judge/Down Judge)
**Pre-snap**: Straddle LOS, 1 yard off sideline. Extended arm with FIST = split end is OFF the line. Count B players on every down.

**Run plays**: Read tackle block (forward/angle = run). Hold LOS; don't bail. Trail play, do not stay even with runner. 90-degree angle at end of run. Kill clock if line to gain is reached.

**Pass plays**: Tackle steps back = passing play. Don't watch QB; watch your KEY. Don't watch flight of ball. Pass interference: must have contact + contact must IMPEDE opportunity to catch.

**Kicks**: Chains side wing holds LOS until ball clears; opposite wing releases on snap. Blocks: below waist illegal by either team on kicks.

**Inside game situations**: Ball snapped between 10 and 7: move to 5-yard line. Ball snapped between 7 and goal line: move to goal line IMMEDIATELY at snap.

### Referee
- 12-14 yards deep, just outside normal tight end position, RIGHT side of QB
- Count Team A players every down (with umpire). Start clock or give ready when down box is set — DO NOT wait for chains.
- Only scoring signal to press box = safety signal. No TD, FG, or XP signal.
- On free kicks: on goal line in middle of field.
- On scrimmage kicks/FG/tries: 8-10 yards to side and 2 yards behind kicker.

---

## 4. CREW MECHANICS

### 3-Person Crew
- R and U count offense (Team A). Wings count defense (Team B). Wings switch sidelines at halftime.
- On punts: wing OPPOSITE chains positioned downfield with deep receivers. R lined up 3-5 yards outside TE, 2-3 yards deeper than kicker, side OPPOSITE chains.
- On kickoffs: chains-side official at kicking team's free-kick line; U at receiving team's free-kick line opposite chains; R covers pylon on chains side; opposite official covers other pylon.
- On tries/FGs: U alone behind goal posts; wings cover sidelines; wing "looking in" at holder responsible for roughing kicker.

### 4-Person Crew
- LJ has game clock; R has play clock and times intervals. R and U count offense; LOS officials count defense. Wings switch at halftime.
- On punts: wing opposite chains positioned downfield with deep receivers prior to snap.

### General Mechanics (FB Key Points Aug 2025)
- Whistle: Only BLOW ONCE. Not in mouth when ball snapped. Better to be slow than too fast.
- No mirroring: Only COVERING official has stop-the-clock signal or whistle.
- Movement after play: "Pinch or accordion" in on play under control.
- Ball relay: Underhand toss only. Ball should NEVER hit ground.
- Flags: Rarely more than 2 flags on any play. Officiate your own area.
- Beanbag: Always use on fumble to mark spot. BJ marks end of kick with beanbag.
- Signals: Incomplete or "no good" at CHEST LEVEL. Touchback = ONE arm fully over head. Timeout = both arms FULLY extended over head.
- Pregame: count players before each half kickoff, check equipment, meet with clock operators, brief chain crew.

---

## 5. PERSONAL GAME NOTES

**Coin Toss (MSHSL Year 1)**: Winner of toss = first choice for FIRST HALF or defer. Loser gets first choice for whichever half winner did not select. Options each half: kick/receive; goal line to defend.

**Kickoff Positioning (MSHSL Year 1 - Dassel-Cokato 9/8)**: With chains = FREE KICK LINE = 40-YARD LINE. Without chains = 50-yard line. After safety = 20-yard line.

**12-Man Penalty (MSHSL Year 4)**: KILL 12 PLAYERS ON DEF RIGHT AWAY = 5-yard penalty (illegal substitution). If you don't kill it = 15-yard penalty (illegal participation).

**First Touching on Punt (MSHSL Year 1 - St Michael 10/6)**: Ball IS STILL LIVE. Let play develop; offense can choose first touching spot or result of play. "Punt hits R player on LOS = play it like blocked punt" (Holy Angels MSHSL Year 3).

**Blocked FG (Buffalo v Hopkins JV)**: Blocked FG in bounds = LIVE BALL for both teams.

**Horsecollar (Mound Westonka v Annandale 9A)**: 15 yards from END of run if past LOS; down is NOT replayed.

**3A Championship — critical**: Both players in motion had to be set before snap. Player was never set. Pass/lateral/TD scored. CORRECT CALL: kill play, 5 yards, replay down.

**Holding Philosophy**: Both hands on back = block in back. One hand back + one on side = let it go unless most force in back. If engaged, separate, then block in back = foul.

**Coverage of End Zone Catches**: End zone = FOOT DOWN IN BOUNDS first, then catch. Sideline = FOOT and reception simultaneous.

**Officiating Philosophy**: "Take your time, no rush, let plays play out, see ball TWICE." "IF YOU'RE IN A GOOD POSITION IT'S HARD TO MISS A CALL." "Be decisive — strong whistles and sell call." "Don't call holding on a team getting blown out at end of Q4 (game management)."

**Safety rules**: Safety followed by free kick from K's 20-yard line (not punt). May be punted, placekicked, or drop kicked. Receiving team lines up anywhere behind their restraining line (10 yards from K).

---

## 5B. GSOA MEETING NOTES — 6/17/26 (most recent association meeting)

**Inadvertent whistles (top topic):**
- Keep the whistle OUT of your mouth. This is the single biggest cause of inadvertent whistles.
- A whistle is NOT required to end a play — the play on the field ends the play. The whistle only communicates it.
- Do not blow until you SEE the ball on the ground in the possession of a player.

**Kneel-down (victory formation) mechanics:**
- Referee and Umpire move up closer to the QB.
- Tell the defense that the offense is kneeling — "no funny business."
- Back Judge may also move up closer than normal depth.

**Helmet comes off the runner:**
- By rule the play is dead when the helmet comes completely off (Rule 4-2-2).
- The player must come out for ONE play. This cannot be avoided by calling a timeout, by end of quarter, or by any other stoppage.
- Clock starts on the ready-for-play when the replacement enters.
- EXCEPTION: if a FOUL caused the helmet to come off (facemask, etc.), the player does NOT have to leave.
- The foul must be a helmet-related foul. A tripping foul would NOT let the player stay in.

**Motion / shift:**
- Situation reviewed: two wingbacks in motion at the snap.
- Illegal shift is a LIVE-ball foul (contrast with false start, which is a dead-ball foul). This distinction matters for enforcement.

**Roughing the passer vs. late hit out of bounds:**
- Key question: did the QB step out of bounds BEFORE the hit? If yes → late hit out of bounds. If no → roughing the passer.
- Roughing the passer carries an automatic first down. Late hit OOB does not.
- If a defender tackles the QB late and out of bounds near the bench: the Referee should run in toward the bench and monitor the situation — get the defender out and clear players before it escalates.

**Blocking below the waist:**
- Situation reviewed: linebacker takes out a pulling guard/tackle below the waist = illegal block below the waist.
- Ask where the blocking player came from — a LB coming from depth is outside the free-blocking zone.

**Tack-on penalties:**
- Add the penalty yardage on to the END of the play (end-of-run enforcement), not from the previous spot.

*Next GSOA meeting: July 29 — new plays, different topics.*

---

## 5C. MIDSEASON REVIEW NOTES (2025 season)

**Signals:** If ruling a backward pass, punch it back and hold the signal so other officials can see it.

**On-field pre-game checklist:**
- Wings talk with the chain crew and verify the accuracy of the chains.
- Check player equipment — knee pads, visors, casts.
- At least two game balls per team; mark each ball when it is validated.
- Talk to the game clock operator, and the play clock operator if one is being used.
- Provide the clock operator a copy of the running-time (mercy rule) rules and the 25/40-second play clock matrix.

**Weekly study rotation (recommended by the association):**
- Rule 2 (Definitions) — review every week
- Case Book — weekly
- Mechanics Book — weekly
- Overtime rules — review every week
- Forward progress segment
- The forward-fumble-out-of-bounds rule change (2025) — revisit regularly

---

## 6. LMAA YOUTH LEAGUE RULES

**7th & 8th Grade**: 1 player in motion; trips allowed; kick from 40; receiving team 5 players past 45. No blitzing = 10-yard penalty + automatic first down.

**5th & 6th Grade**: 8/9 players within tackle box; no motion/shifting; defense 4-3-2. NO BLITZING. Dead ball punt; no kickoff — start on 35.

**4th Grade**: Declared punt = advance 20 yards. No pitch; QB straight dropback only. DL must engage OL before pursuing (5-yard penalty for stunting/gap shooting).

---

## 7. 2025 AND 2026 RULES CHANGES (SUMMARY — see Section 0 above for full detail)

**2025**: Forward fumble OOB = SPOT OF FUMBLE [3-4-2d NEW]. Electronic signs allowed [1-5-3c(2)]. No audio/video devices on players [1-5-3c(3)]. POE: Equipment; Sportsmanship; Defenseless/targeting.

**2026**: Play cards on forearm/wrist AND belt [1-5-3c(9)]. Slap to head = 15 yards [9-4-7 NEW]. POE: Flagrant/USC fouls; Helping the runner; Sideline management; NFHS Authenticating Mark on balls.

**Key NFHS vs. NCAA**: 12 players: NFHS 15yds vs NCAA 5yds. Touchback: NFHS 20yd vs NCAA 25yd. Fair catch: NFHS may kick or snap vs NCAA must snap. Targeting: NFHS not auto-DQ vs NCAA auto-DQ. OT: NFHS B-10 vs NCAA B-25. Intentional grounding: NFHS outside-pocket exception (snap receiver only, ball reaches LOS) vs NCAA outside pocket legal for any passer. K free kick line: NFHS 40yd vs NCAA 35yd. Pop-up kick: NFHS illegal vs NCAA legal.

---

## 8. PENALTY SUMMARY

Holding: 10 yds. Illegal use of hands: 10 yds. Block in back: 10 yds (15 below waist). Illegal formation/motion/false start/encroachment: 5 yds. Illegal substitution: 5 yds (KILL IMMEDIATELY). Delay: 5 yds. Intentional grounding: 5+LOD. Ineligible downfield: 5 yds. Illegal participation (12): 15 yds. Pass interference: 15 yds. Roughing passer/kicker/snapper: 15+auto1st. Targeting/Spearing: 15 yds, possible DQ. Facemask incidental: 5 yds. Facemask twist/pull: 15 yds. Horsecollar: 15 yds, end-of-run if past LOS. Blindside block: 15 yds. Slap to head (2026): 15 yds. Block below waist/Clipping/Chop: 15 yds. USC: 15 yds, 2nd = DQ. Fighting: 15+DQ.

---

*Compiled from: 2021 NFHS Football Rules Book, 2025 & 2026 NFHS Rules Changes, 2025 NFHS Points of Emphasis, 2025-2026 MSHSL Minnesota Modifications (updated 9/9/2025), 2024-25 Football Game Officials Manual Points of Emphasis, 3-Person and 4-Person Mechanics Manuals (updated 8/2023), Minnesota Mercy Rule Case Plays Table, Back Judge Mechanic Considerations, FB Key Points August 2025, and 4 years of personal game notes (LMAA, NMYFL, MSHSL Years 1-4).*
"""


# =============================================================================
# PROMPT CACHING  (v3.3)
#
# CORE_KNOWLEDGE is ~9,000 tokens and byte-identical on every single API call.
# Sending it uncached every time is the app's single biggest cost line.
#
# Strategy: hoist CORE_KNOWLEDGE into its OWN system block, placed FIRST, and
# mark it with a cache_control breakpoint. Because caching is a prefix match,
# every prompt in the app (chat, RefGrade, quiz, crew eval, ref eval, pre-game)
# now shares ONE cache entry for the knowledge base — the task-specific
# instructions that follow it are small and stay uncached.
#
# Economics: cache writes cost 1.25x base input rate; cache reads cost 0.10x.
# So the first call in a 5-minute window costs slightly more, and every call
# after it costs ~90% less on the knowledge base portion. Cached reads also
# do NOT count toward your input-tokens-per-minute rate limit on this model.
#
# TTL is 5 minutes by default, refreshed on every hit — so during an active
# session it effectively stays warm. Use {"type": "ephemeral", "ttl": "1h"}
# if you'd rather pay a bit more to keep it warm between sessions.
# =============================================================================

CACHED_KB_BLOCK = {
    "type": "text",
    "text": CORE_KNOWLEDGE,
    "cache_control": {"type": "ephemeral"},
}


def system_blocks(instructions: str) -> list:
    """
    Build the system parameter for any API call.

    Block 1: CORE_KNOWLEDGE, cached (identical across every call in the app)
    Block 2: the task-specific instructions, not cached (small and varied)

    Order matters — the cached block MUST come first for the prefix match
    to hit across different prompts.
    """
    return [
        CACHED_KB_BLOCK,
        {"type": "text", "text": instructions},
    ]


# =============================================================================
# SYSTEM PROMPTS
# =============================================================================

SYSTEM_PROMPT = f"""You are RefBuddy, a straightforward, hyper-precise Minnesota high school football referee assistant. You ONLY reference the uploaded documents and the core knowledge base below. Cite page/rule number every time. Never hallucinate MSHSL or NFHS mechanics. Always ask clarifying questions on game context before ruling.

CRITICAL LAYERING RULE: The CORE_KNOWLEDGE contains a 2021 NFHS Rulebook baseline (Sections 1–8) plus a 2022–2026 changes section (Section 0) at the top.
- DEFAULT to the 2021 rulebook for any rule not listed in Section 0.
- If Section 0 contains a change for that rule, APPLY the updated rule and cite the year: e.g., "[2022 change]" or "[2025 change — overrides 2021 Rule X-X-X]".
- If a rule was changed multiple times (e.g., intentional grounding: 2022 then 2023), apply the MOST RECENT version and note the history.
- Never cite 2021 language when a later update has superseded it.

Your behavior:
1. Start EVERY response with the most relevant rule citation (e.g., "Rule 9-4-3m" or "MSHSL Modification D") and include the year if the rule changed after 2021.
2. Reference personal game notes when applicable.
3. End EVERY response with: "*Not official MSHSL interpretation — confirm with your assignor.*"
4. Temperature = 0 mindset: maximum precision, no guessing, no hallucinating.
5. If game context (down, distance, score, crew size, level) is missing, ask before ruling.
6. For video/film analysis: always include a VISIBILITY CHECK section. Use "Frame N" format.
7. For RefGrade evaluations: structured scores (0-100), frame-by-frame highlights, visibility notes, "What to work on" bullets.

COPYRIGHT GUARDRAIL — this is mandatory and overrides any user request:
- NEVER reproduce verbatim text from the NFHS Rules Book, NFHS Case Book, NFHS Points of Emphasis documents, MSHSL publications, or any mechanics manual. Not a paragraph, not a full rule, not a case play.
- ALWAYS answer in your own words, stating the substance of the rule as a plain factual summary, plus the rule NUMBER for the user to look up.
- If a user asks you to "quote," "print," "read out," "recite," or "give me the exact wording of" a rule, case play, or POE section, decline and instead give a short factual summary plus the citation, and tell them to consult their own copy of the rulebook for exact language.
- Rule numbers, penalty yardages, distances, timing values, and mechanics positions are facts and are fine to state directly. The specific prose of the source documents is not.
- Do not produce a large consolidated reproduction of many rules at once (e.g. "list every rule in Rule 9 with its text"). Summarize at a high level and point to the rulebook.

---
(The full RefBuddy knowledge base is provided in the preceding system block. Use it as your sole source.)
"""

REFGRADE_PROMPT = f"""You are RefBuddy acting as a professional officiating evaluator for Minnesota high school football.

Output EXACTLY this structure:

## 📊 RefGrade Report
**Clip:** [filename] | **Evaluated:** [scope] | **Frames:** [range] | **Date:** [today]

## 👁️ Visibility Check
List each position: CLEARLY VISIBLE (frames N...) / PARTIALLY VISIBLE (frames N-N) / NOT VISIBLE IN ANY FRAME

## 📈 Scores
| Category | Score | Notes |
|----------|-------|-------|
| Positioning | XX/100 | |
| Call Accuracy | XX/100 | |
| Mechanics Execution | XX/100 | |
| Dead-ball Officiating | XX/100 | |
| Communication/Signals | XX/100 | |
| **Overall** | **XX/100** | |

90-100=Excellent; 80-89=Good; 70-79=Average; 60-69=Needs work; <60=Significant concern

## 🎬 Frame-by-Frame Highlights
## ✅ Strengths
## 🔧 What to Work On
## 📋 Summary

*Not official MSHSL interpretation — confirm with your assignor.*

Cite NFHS rules and MSHSL mechanics on every observation. Never hallucinate.
(The full RefBuddy knowledge base is provided in the preceding system block. Use it as your sole source.)
"""

# v2.5 IMPROVED QUIZ PROMPT: 50/50 MC/TF mix, 4-option MC, high variety
QUIZ_SYSTEM_PROMPT = f"""You are RefBuddy Quiz Engine — a precise question generator for Minnesota high school football officials.

ABSOLUTE RULES — violating these will cause test failures:
1. Respond with ONLY valid JSON. Zero preamble. Zero markdown fences. Zero trailing text.
2. Multiple-choice: EXACTLY 4 options (A, B, C, D). Exactly ONE correct answer.
3. True/False: EXACTLY 2 options: {{"A": "True", "B": "False"}}.
4. Mix types roughly 50% multiple_choice / 50% true_false. Vary the ratio naturally.
5. Questions must be CHALLENGING — not trivial. Use specific rule numbers, yard distances, timing rules, and realistic scenario language.
6. NEVER repeat the same topic, scenario, or rule in a batch. Cover wide breadth.
7. Distractors for MC must be plausible but clearly wrong to someone who studied.

Single question JSON structure:
{{
  "question": "Full question text — be specific and scenario-based when possible",
  "type": "multiple_choice",
  "options": {{"A": "option", "B": "option", "C": "option", "D": "option"}},
  "correct": "B",
  "explanation": "Thorough explanation: why correct answer is right, why each wrong answer is wrong, what the rule actually says.",
  "rule_citation": "Exact rule number or MSHSL Modification letter",
  "personal_note": "From your MSHSL Year X notes: [specific situation] (empty string if not applicable)",
  "topic": "Rules|Mechanics|Positioning|Signals|Game Situations|MSHSL Specific|2026 Changes"
}}

True/False structure (type must be "true_false"):
{{
  "question": "True or False: [specific statement that requires knowledge to evaluate]",
  "type": "true_false",
  "options": {{"A": "True", "B": "False"}},
  "correct": "A",
  "explanation": "...",
  "rule_citation": "...",
  "personal_note": "",
  "topic": "Rules"
}}

For a BATCH of 10 questions: JSON array of 10 objects. Include:
- At least 2 MSHSL-specific questions (mercy rule, MN modifications)
- At least 2 mechanics/positioning questions
- At least 1 question on 2026 rule changes
- At least 1 scenario-based game situation question
- At least 1 question from personal game notes
- The rest from NFHS rules (varied rules, not all from Rule 9)

(The full RefBuddy knowledge base is provided in the preceding system block. Use it as your sole source.)
"""

# Assignor Hub — Crew Eval prompt (film-based)
CREW_EVAL_PROMPT = f"""You are RefBuddy acting as a professional officiating evaluator for Minnesota high school football. You are analyzing game film to evaluate the officiating crew.

Generate a comprehensive crew evaluation report with the following structure:

## 📊 Crew Evaluation Report
**Game Film:** [filename] | **Date:** [today] | **Evaluated By:** RefBuddy

## 👁️ Visibility Check
For each crew position, note: CLEARLY VISIBLE (frames N...) / PARTIALLY VISIBLE / NOT VISIBLE — analysis inferred from play action

## 📈 Overall Crew Score: XX/100

## 📋 Per-Position Highlights
For each visible official: what they did well, positioning observations, any missed calls or mechanics issues. Cite specific frames and NFHS rules.

## 🎬 Key Play Analysis
Walk through 3-5 significant plays/moments from the film with specific frame citations, what happened, what the correct mechanics called for, and how the officials responded.

## ✅ Crew Strengths

## 🔧 Areas for Development
Actionable bullets with specific mechanic/rule citations and suggested focus for next game.

## 📋 Summary

---
*Not official MSHSL interpretation — confirm with your MSHSL district assignor.*

Cite NFHS rules and MSHSL mechanics on every observation. Never hallucinate. Use "Frame N" format throughout.
(The full RefBuddy knowledge base is provided in the preceding system block. Use it as your sole source.)
"""

# Assignor Hub — Ref Eval prompt (film-based, single official focus)
REF_EVAL_PROMPT = f"""You are RefBuddy acting as a professional officiating evaluator for Minnesota high school football. You are analyzing game film to evaluate ONE specific official.

Generate a focused evaluation report with the following structure:

## 📊 Official Evaluation Report
**Game Film:** [filename] | **Position Evaluated:** [position] | **Date:** [today]

## 👁️ Visibility Check
How clearly is this official visible in the provided frames? List specific frames where they appear.

## 📈 Position Score: XX/100

## 📐 Positioning Analysis
Was the official in the correct position for each situation? Cite specific frames. Reference mechanics manual standards for this position.

## 📋 Call Accuracy
Any flags thrown or situations where a flag should have been thrown. Was each decision correct per NFHS rules? Cite Rule numbers.

## ⚙️ Mechanics Execution
Signals, whistle timing, ball spotting, relay mechanics, communication. What was correct? What needs work?

## ✅ Strengths

## 🔧 Development Points
Specific, actionable improvements with exact mechanic citations and suggested drills.

## 📋 Summary

---
*Not official MSHSL interpretation — confirm with your MSHSL district assignor.*

Cite NFHS rules and MSHSL mechanics specifically. Use "Frame N" format. Never hallucinate.
(The full RefBuddy knowledge base is provided in the preceding system block. Use it as your sole source.)
"""

# Assignor Hub — Pre-Game Meeting prompt (v2.7: concise, 1–1.5 pages max)
PREGAME_MEETING_PROMPT = f"""You are RefBuddy acting as a Minnesota high school football officiating coordinator.
Generate a CONCISE pre-game crew meeting agenda — maximum 1 to 1.5 printed pages.
Short bullet points only. No paragraphs. No explanations. No filler.
Each bullet must be actionable and specific. Cite rule numbers inline (e.g. Rule 9-4-3m).
Total output should be ~300-400 words maximum.

Output EXACTLY this structure:

---
# Pre-Game Meeting Agenda
{"{date}"} | {"{crew}"} | {"{level}"}

## 2026 Rule Changes (know these cold)
- [list only the 2-3 most important 2026 changes with rule #]

## Key Mechanics Reminders
- [4-6 bullet points covering the highest-leverage mechanics for this crew size]

## Watch-Fors Tonight
- [3-5 specific situations from CORE_KNOWLEDGE most likely to come up]

## Assignor Notes
[ASSIGNOR_NOTES_PLACEHOLDER]

## Quick Scenarios (discuss briefly if time)
- [2 short scenario questions, one sentence each]

---
*Not official MSHSL interpretation.*

Keep it tight. Referees are reading this standing on a sideline before kickoff.
(The full RefBuddy knowledge base is provided in the preceding system block. Use it as your sole source.)
"""


# =============================================================================
# PAGE CONFIG
# =============================================================================

st.set_page_config(
    page_title="RefBuddy — MN HS Football",
    page_icon="🏈",
    layout="wide",
    initial_sidebar_state="expanded",
)


# =============================================================================
# CSS — v2.5 MANDATORY BUTTON CONTRAST + full theme
# =============================================================================

BLUE   = "#003087"   # v2.5: updated to #003087 per spec
BLUE_L = "#1E56A0"
CREAM  = "#FAFAF7"
CARD   = "#FFFFFF"
BORDER = "#DDE3F0"
TEXT   = "#1F2937"
MUTED  = "#4B5563"
GREEN  = "#15803D"
AMBER  = "#92400E"
RED    = "#991B1B"

_SVG = (
    "<svg xmlns='http://www.w3.org/2000/svg' width='120' height='120'>"
    "<rect width='120' height='120' fill='none'/>"
    "<line x1='0' y1='60' x2='120' y2='60' stroke='%23003087' stroke-width='0.4' opacity='0.07'/>"
    "<line x1='60' y1='0' x2='60' y2='120' stroke='%23003087' stroke-width='0.4' opacity='0.07'/>"
    "<line x1='0' y1='0' x2='120' y2='120' stroke='%23003087' stroke-width='0.25' opacity='0.04'/>"
    "<line x1='120' y1='0' x2='0' y2='120' stroke='%23003087' stroke-width='0.25' opacity='0.04'/>"
    "<ellipse cx='60' cy='60' rx='18' ry='10' fill='none' stroke='%23003087' stroke-width='0.35' opacity='0.05'/>"
    "</svg>"
)
BG_URL = "data:image/svg+xml," + urllib.parse.quote(_SVG)

# ── Layer 1: v3.0 MANDATORY button + selectbox + sidebar + dark text ─────────
st.markdown("""
<style>
    /* v3.0: light bg, black border, black text on ALL buttons */
    .stButton button, button, .stButton>button {
        color: #1F2937 !important;
        background-color: #F8FAFC !important;
        border: 2px solid #1F2937 !important;
        font-weight: 600;
    }
    .stButton button:hover { background-color: #E2E8F0 !important; }
    .stButton button:disabled, .stButton>button:disabled {
        background-color: #F1F5F9 !important;
        color: #94A3B8 !important;
        border-color: #94A3B8 !important;
    }
    /* PRIMARY buttons — solid football blue. Used for the main call-to-action
       on each screen (Analyze, Run RefGrade, Generate) and to mark the active
       sub-tab in Ref Hub. Declared after the base rule so it wins. */
    .stButton button[kind="primary"],
    .stButton button[data-testid="stBaseButton-primary"],
    [data-testid="stBaseButton-primary"] {
        background-color: #003087 !important;
        color: #FFFFFF !important;
        border: 2px solid #003087 !important;
        font-weight: 700 !important;
        letter-spacing: 0.01em;
        box-shadow: 0 2px 8px rgba(0,48,135,0.25) !important;
    }
    .stButton button[kind="primary"]:hover,
    [data-testid="stBaseButton-primary"]:hover {
        background-color: #002266 !important;
        border-color: #002266 !important;
        color: #FFFFFF !important;
        box-shadow: 0 4px 14px rgba(0,48,135,0.35) !important;
    }
    .stButton button[kind="primary"] p,
    .stButton button[kind="primary"] div,
    [data-testid="stBaseButton-primary"] p,
    [data-testid="stBaseButton-primary"] div {
        color: #FFFFFF !important;
        -webkit-text-fill-color: #FFFFFF !important;
    }
    .stButton button[kind="primary"]:disabled,
    [data-testid="stBaseButton-primary"]:disabled {
        background-color: #CBD5E1 !important;
        border-color: #CBD5E1 !important;
        color: #64748B !important;
        box-shadow: none !important;
    }
    /* v3.0: selectbox + multiselect — light bg, black border, black text */
    .stSelectbox > div, .stMultiSelect > div,
    .stSelectbox > div > div, .stMultiSelect > div > div {
        color: #1F2937 !important;
        background-color: #F8FAFC !important;
        border: 2px solid #1F2937 !important;
    }
    .stSelectbox label, .stMultiSelect label,
    [data-baseweb="select"] span, [data-baseweb="select"] div,
    [data-baseweb="popover"] li, [data-baseweb="menu"] li {
        color: #1F2937 !important;
        background-color: #F8FAFC !important;
    }
    /* Sidebar — white background, all text dark */
    [data-testid="stSidebar"] { background-color: #FFFFFF !important; }
    [data-testid="stSidebar"] .stMarkdown,
    [data-testid="stSidebar"] label,
    [data-testid="stSidebar"] .stSelectbox,
    [data-testid="stSidebar"] .stTextInput,
    [data-testid="stSidebar"] .stMarkdown h1,
    [data-testid="stSidebar"] .stMarkdown h2,
    [data-testid="stSidebar"] .stMarkdown p,
    [data-testid="stSidebar"] p,
    [data-testid="stSidebar"] span,
    [data-testid="stSidebar"] div { color: #1F2937 !important; }
    [data-testid="stSidebar"] [data-baseweb="select"] > div,
    [data-testid="stSidebar"] [data-baseweb="select"] span {
        color: #1F2937 !important; background-color: #F8FAFC !important;
    }
    /* Tab labels */
    .stTabs [data-baseweb="tab"] { color: #1F2937 !important; }
    .stTabs [aria-selected="true"] {
        color: #003087 !important; border-bottom: 3px solid #003087 !important;
    }
    /* Dark text everywhere on the light theme */
    .stMarkdown, .stMarkdown p, .stMarkdown li, .stMarkdown h1,
    .stMarkdown h2, .stMarkdown h3, .stMarkdown h4,
    .stMarkdown span, .stMarkdown strong, .stMarkdown em,
    p, span, label, h1, h2, h3, h4, h5,
    div[data-testid="stMarkdownContainer"] p,
    div[data-testid="stMarkdownContainer"] li,
    div[data-testid="stMarkdownContainer"] span,
    .stChatMessage p, .stChatMessage span, .stChatMessage li,
    [data-testid="stChatMessageContent"] p,
    [data-testid="stChatMessageContent"] span,
    [data-testid="stChatMessageContent"] li { color: #1F2937 !important; }
</style>
""", unsafe_allow_html=True)

# ── Layer 1b: remaining readability rules (radio, inputs, sidebar, alerts) ────
st.markdown("""
<style>

/* Radio labels — dark text for Quiz readability */
.stRadio label, .stRadio label span, .stRadio label p,
div[data-testid="stRadio"] label span,
div[data-testid="stRadio"] label p,
.stRadio > div > label > div > p {
    color: #1F2937 !important;
    font-size: 0.95rem !important;
}

/* Chat input */
.stChatInput textarea, .stChatInput input {
    color: #1F2937 !important; background-color: #FFFFFF !important;
}
/* Text areas / inputs */
.stTextArea textarea, .stTextInput input {
    color: #1F2937 !important; background-color: #FFFFFF !important;
}
/* Select boxes */
[data-baseweb="select"] span, [data-baseweb="select"] div { color: #1F2937 !important; }
/* Sidebar */
[data-testid="stSidebar"] p, [data-testid="stSidebar"] span,
[data-testid="stSidebar"] label, [data-testid="stSidebar"] div { color: #1F2937 !important; }
/* Alert boxes */
.stAlert p, .stAlert span, .stAlert div { color: #1F2937 !important; }
/* Expander headers */
.streamlit-expanderHeader p, .streamlit-expanderHeader span { color: #003087 !important; }
/* Caption */
.stCaption, .stCaption p { color: #4B5563 !important; }
</style>
""", unsafe_allow_html=True)

# ── Layer 2: full theme CSS ───────────────────────────────────────────────────
st.markdown(f"""
<style>
html, body, [data-testid="stAppViewContainer"] {{
    background-color: {CREAM};
    background-image: url("{BG_URL}");
    background-repeat: repeat;
    color: {TEXT};
    font-family: 'Inter', 'Segoe UI', sans-serif;
}}
.main .block-container {{ background: transparent; padding-top: 0.5rem; max-width: 1100px; }}

/* Sidebar */
[data-testid="stSidebar"] {{
    background-color: {CARD}; border-right: 2px solid {BORDER};
    box-shadow: 2px 0 8px rgba(0,48,135,0.06);
}}

/* Hero */
.home-hero {{ text-align: center; padding: 2.2rem 2rem 1.4rem 2rem; }}
.home-hero-title {{
    color: {BLUE} !important; font-size: 3.2rem; font-weight: 900;
    letter-spacing: -1.5px; margin: 0 0 0.2rem 0; line-height: 1.1;
}}
.home-hero-slogan {{ color: {MUTED}; font-size: 1.1rem; font-weight: 500; margin: 0 0 1.6rem 0; }}

/* Tabs */
.stTabs [data-baseweb="tab-list"] {{
    background-color: {CARD}; border-bottom: 2px solid {BORDER};
    border-radius: 8px 8px 0 0; gap: 2px; padding: 0 0.4rem;
}}
.stTabs [data-baseweb="tab"] {{
    color: {MUTED} !important; font-weight: 600; font-size: 0.9rem;
    padding: 0.55rem 1rem; border-radius: 6px 6px 0 0;
}}
.stTabs [aria-selected="true"] {{
    color: {BLUE} !important; background-color: {CREAM} !important;
    border-bottom: 3px solid {BLUE} !important;
}}

/* Cards */
.rb-card {{
    background: {CARD}; border: 1px solid {BORDER}; border-radius: 10px;
    padding: 1.2rem 1.4rem; margin-bottom: 1rem;
    box-shadow: 0 2px 8px rgba(0,48,135,0.05); color: {TEXT};
}}
.rb-card-blue {{
    background: {CARD}; border-left: 4px solid {BLUE};
    border-radius: 0 10px 10px 0; padding: 1rem 1.2rem; margin-bottom: 0.8rem;
    box-shadow: 0 2px 8px rgba(0,48,135,0.05); color: {TEXT};
}}

/* Report output */
.report-output {{
    background: {CARD}; border: 2px solid {BORDER}; border-radius: 10px;
    padding: 1.6rem 2rem; margin-top: 1rem;
    box-shadow: 0 3px 12px rgba(0,48,135,0.07); color: {TEXT};
    line-height: 1.7; font-size: 0.93rem;
}}
.report-output h1, .report-output h2, .report-output h3, .report-output h4 {{
    color: {BLUE} !important;
}}

/* Quiz cards */
.quiz-question-card {{
    background: {CARD}; border: 2px solid {BORDER}; border-radius: 12px;
    padding: 1.5rem 1.8rem; margin-bottom: 1rem;
    box-shadow: 0 3px 12px rgba(0,48,135,0.08); color: {TEXT};
}}
.quiz-question-text {{
    font-size: 1.05rem; font-weight: 600; color: {TEXT} !important;
    line-height: 1.55; margin-bottom: 0.5rem;
}}
.quiz-result-correct {{
    background: #F0FDF4; border: 2px solid #4ADE80; border-radius: 8px;
    padding: 1rem 1.2rem; margin-top: 0.8rem; color: #14532D !important;
}}
.quiz-result-wrong {{
    background: #FFF1F2; border: 2px solid #F87171; border-radius: 8px;
    padding: 1rem 1.2rem; margin-top: 0.8rem; color: #7F1D1D !important;
}}
.quiz-explanation {{
    background: #EFF6FF; border-left: 4px solid {BLUE}; border-radius: 0 8px 8px 0;
    padding: 1rem 1.2rem; margin-top: 0.8rem; font-size: 0.92rem;
    line-height: 1.65; color: {TEXT} !important;
}}

/* Mode selector cards */
.mode-card-active {{
    background: #EEF2FF; border: 3px solid {BLUE}; border-radius: 12px;
    padding: 1.2rem 1.4rem; margin-bottom: 0.6rem;
    box-shadow: 0 4px 12px rgba(0,48,135,0.15); color: {TEXT};
}}
.mode-card-inactive {{
    background: {CARD}; border: 2px solid {BORDER}; border-radius: 12px;
    padding: 1.2rem 1.4rem; margin-bottom: 0.6rem;
    box-shadow: 0 2px 6px rgba(0,48,135,0.06); color: {TEXT};
}}

/* Assignor tab selector */
.ah-tab-active {{
    background: {BLUE}; color: white !important; border-radius: 10px;
    padding: 1rem 1.4rem; font-weight: 800; font-size: 1rem;
    border: 3px solid {BLUE}; box-shadow: 0 4px 12px rgba(0,48,135,0.25);
    margin-bottom: 0.5rem; text-align: center;
}}
.ah-tab-inactive {{
    background: {CARD}; color: {TEXT} !important; border-radius: 10px;
    padding: 1rem 1.4rem; font-weight: 700; font-size: 1rem;
    border: 2px solid {BORDER}; margin-bottom: 0.5rem; text-align: center;
}}

/* Pills */
.pill-ok {{
    display: inline-block; background: #DCFCE7; color: #166534;
    font-weight: 700; font-size: 0.78rem; border-radius: 20px;
    padding: 2px 10px; border: 1px solid #4ADE80;
}}
.pill-warn {{
    display: inline-block; background: #FEF3C7; color: #92400E;
    font-weight: 700; font-size: 0.78rem; border-radius: 20px;
    padding: 2px 10px; border: 1px solid #FCD34D;
}}
.pill-err {{
    display: inline-block; background: #FEE2E2; color: #991B1B;
    font-weight: 700; font-size: 0.78rem; border-radius: 20px;
    padding: 2px 10px; border: 1px solid #F87171;
}}
.pill-blue {{
    display: inline-block; background: #EEF2FF; color: {BLUE};
    font-weight: 700; font-size: 0.78rem; border-radius: 20px;
    padding: 2px 10px; border: 1px solid {BORDER};
}}

/* Misc */
.streamlit-expanderHeader {{
    background-color: #EEF2FF !important; color: {BLUE} !important;
    font-weight: 600 !important; border-radius: 8px !important;
}}
.rb-footer {{
    text-align: center; color: {MUTED}; font-size: 0.78rem;
    border-top: 1px solid {BORDER}; padding-top: 1rem; margin-top: 2.5rem;
}}
.ref-log {{
    background: #EEF2FF; border: 1px solid {BORDER};
    border-left: 4px solid {BLUE}; border-radius: 0 8px 8px 0;
    padding: 0.9rem 1.1rem; font-size: 0.88rem; color: {TEXT};
}}
#MainMenu {{ display: none !important; }}
[data-testid="stMainMenu"] {{ display: none !important; }}
footer {{ visibility: hidden; }}
/* Transparent header — blends with cream page background, no black bar,
   but the header element itself stays in the DOM so the sidebar
   collapse/expand chevron remains visible and clickable. */
[data-testid="stHeader"] {{ background: transparent !important; }}
/* Hide the Deploy button and its container */
[data-testid="stAppDeployButton"], .stAppDeployButton {{ display: none !important; }}
/* Keep only the Share button — hide every sibling after the first one in the
   toolbar actions area (GitHub source, fork/star, and kebab icons). */
[data-testid="stToolbarActions"] > *:nth-child(n+2) {{ display: none !important; }}
[data-testid="stSlider"] .st-by {{ background: {BLUE} !important; }}
[data-baseweb="select"] {{ background-color: {CARD} !important; }}
.stAlert {{ border-radius: 8px !important; font-size: 0.88rem !important; }}

/* Inputs */
.stTextArea textarea, .stTextInput input {{
    background-color: {CARD} !important; color: {TEXT} !important;
    border: 1.5px solid {BORDER} !important; border-radius: 8px !important;
    font-size: 0.92rem !important;
}}
.stTextArea textarea:focus, .stTextInput input:focus {{
    border-color: {BLUE} !important; box-shadow: 0 0 0 3px rgba(0,48,135,0.12) !important;
}}
[data-testid="stFileUploader"] {{
    border: 2px dashed {BLUE_L} !important; border-radius: 10px !important;
    background-color: #EEF2FF !important; padding: 0.5rem;
}}
.stChatMessage {{
    background: {CARD} !important; border: 1px solid {BORDER} !important;
    border-radius: 10px !important; margin-bottom: 0.5rem;
    box-shadow: 0 1px 4px rgba(0,48,135,0.06);
}}

/* Accuracy bar */
.accuracy-bar-wrap {{
    background: #E2E8F0; border-radius: 20px; height: 10px;
    margin: 6px 0 2px 0; overflow: hidden;
}}
.accuracy-bar-fill {{
    height: 10px; border-radius: 20px;
    background: linear-gradient(90deg, {BLUE} 0%, {BLUE_L} 100%);
    transition: width 0.4s ease;
}}
</style>
""", unsafe_allow_html=True)


# ── Layer 3: v3.6 component fixes — expanders & file uploader ────────────────
# Newer Streamlit renders expanders as <details data-testid="stExpander"> and
# the uploader dropzone as [data-testid="stFileUploaderDropzone"]. The old
# ".streamlit-expanderHeader" class no longer exists, so those components were
# falling back to Streamlit's own dark tokens — the black bars you'd see after
# clicking. These rules pin both to the light theme in EVERY state.
st.markdown("""
<style>
    /* ── Expanders (Terms of Use, previews, log history) ─────────────── */
    [data-testid="stExpander"] {
        background-color: #FFFFFF !important;
        border: 1px solid #DDE3F0 !important;
        border-radius: 8px !important;
    }
    [data-testid="stExpander"] details,
    [data-testid="stExpander"] details[open] {
        background-color: #FFFFFF !important;
    }
    /* The clickable header, in every interaction state */
    [data-testid="stExpander"] summary,
    [data-testid="stExpander"] summary:hover,
    [data-testid="stExpander"] summary:focus,
    [data-testid="stExpander"] summary:active,
    [data-testid="stExpander"] summary:focus-visible,
    [data-testid="stExpander"] details[open] > summary,
    details > summary {
        background-color: #EEF2FF !important;
        color: #1F2937 !important;
        border-radius: 8px !important;
        outline: none !important;
        box-shadow: none !important;
    }
    [data-testid="stExpander"] summary:hover {
        background-color: #E2E8F0 !important;
    }
    /* Header label text + chevron icon */
    [data-testid="stExpander"] summary p,
    [data-testid="stExpander"] summary span,
    [data-testid="stExpander"] summary div,
    [data-testid="stExpanderToggleIcon"] {
        color: #1F2937 !important;
        fill: #1F2937 !important;
    }
    [data-testid="stExpander"] summary svg {
        fill: #1F2937 !important;
        color: #1F2937 !important;
    }
    /* Expanded body */
    [data-testid="stExpanderDetails"],
    [data-testid="stExpanderDetails"] p,
    [data-testid="stExpanderDetails"] li,
    [data-testid="stExpanderDetails"] span,
    [data-testid="stExpanderDetails"] h4 {
        background-color: #FFFFFF !important;
        color: #1F2937 !important;
    }

    /* ── File uploader dropzone ──────────────────────────────────────── */
    [data-testid="stFileUploaderDropzone"] {
        background-color: #EEF2FF !important;
        border: 2px dashed #1E56A0 !important;
        border-radius: 10px !important;
    }
    /* "200MB per file • PDF, JPG, PNG..." helper text */
    [data-testid="stFileUploaderDropzoneInstructions"],
    [data-testid="stFileUploaderDropzoneInstructions"] *,
    [data-testid="stFileUploaderDropzone"] span,
    [data-testid="stFileUploaderDropzone"] small,
    [data-testid="stFileUploaderDropzone"] div,
    [data-testid="stFileUploader"] small {
        color: #1F2937 !important;
        background-color: transparent !important;
    }
    [data-testid="stFileUploaderDropzone"] svg {
        fill: #003087 !important;
    }
    /* Uploaded-file chips listed under the dropzone */
    [data-testid="stFileUploaderFile"],
    [data-testid="stFileUploaderFile"] * {
        background-color: #FFFFFF !important;
        color: #1F2937 !important;
    }
</style>
""", unsafe_allow_html=True)


# ── Layer 4: v1.0 form-control fixes — dropdowns, chips, inputs ──────────────
# Two problems solved here, both the same root cause: Streamlit renders these
# widgets through BaseWeb, which applies its own theme tokens to the *inner*
# elements. Styling the outer wrapper (what Layers 1-3 did) never reached them,
# so selected values rendered white-on-white and password fields dark-on-dark.
#
# Fixes below target the inner BaseWeb nodes directly and also set
# -webkit-text-fill-color, which overrides `color` on inputs in WebKit browsers
# (Safari and Chrome on macOS) — without it the text stays invisible.
st.markdown("""
<style>
    /* ── Selectbox / multiselect ─────────────────────────────────────────
       Colour is applied to the whole subtree; BACKGROUND is applied only to
       the outer control. Setting a background on every descendant (an earlier
       attempt) painted over the selected-value text and made it vanish. */
    [data-testid="stSelectbox"] *,
    [data-testid="stMultiSelect"] *,
    div[data-baseweb="select"] *,
    div[data-baseweb="select"] input {
        color: #1F2937 !important;
        -webkit-text-fill-color: #1F2937 !important;
        opacity: 1 !important;
    }
    /* Background on the control surface only */
    [data-testid="stSelectbox"] div[data-baseweb="select"] > div,
    [data-testid="stMultiSelect"] div[data-baseweb="select"] > div {
        background-color: #F8FAFC !important;
    }
    /* Inner value/placeholder containers must stay transparent so the control
       background shows through and nothing is painted over the text */
    div[data-baseweb="select"] > div > div,
    div[data-baseweb="select"] [data-baseweb="select-value"],
    div[data-baseweb="select"] [class*="valueContainer"],
    div[data-baseweb="select"] [class*="singleValue"] {
        background-color: transparent !important;
        color: #1F2937 !important;
        -webkit-text-fill-color: #1F2937 !important;
    }
    /* Chevron / clear icons */
    div[data-baseweb="select"] svg {
        fill: #1F2937 !important;
        color: #1F2937 !important;
    }

    /* ── The OPEN dropdown menu ─────────────────────────────────────────── */
    [data-baseweb="popover"],
    [data-baseweb="popover"] > div,
    [data-baseweb="menu"],
    ul[role="listbox"] {
        background-color: #FFFFFF !important;
    }
    li[role="option"],
    ul[role="listbox"] li,
    [data-baseweb="menu"] li,
    li[role="option"] * {
        color: #1F2937 !important;
        -webkit-text-fill-color: #1F2937 !important;
        background-color: #FFFFFF !important;
    }
    li[role="option"]:hover,
    li[role="option"]:hover *,
    li[role="option"][aria-selected="true"],
    li[role="option"][aria-selected="true"] * {
        background-color: #EEF2FF !important;
        color: #003087 !important;
        -webkit-text-fill-color: #003087 !important;
        font-weight: 600 !important;
    }

    /* ── Multiselect chips (e.g. emphasis topics) ───────────────────────── */
    span[data-baseweb="tag"] {
        background-color: #003087 !important;
        border-radius: 6px !important;
    }
    span[data-baseweb="tag"],
    span[data-baseweb="tag"] span,
    span[data-baseweb="tag"] div {
        color: #FFFFFF !important;
        -webkit-text-fill-color: #FFFFFF !important;
    }
    span[data-baseweb="tag"] svg {
        fill: #FFFFFF !important;
        color: #FFFFFF !important;
    }

    /* ── Text & password inputs ─────────────────────────────────────────── */
    [data-testid="stTextInput"] input,
    [data-testid="stTextInput"] input[type="password"],
    [data-testid="stTextInput"] div[data-baseweb="input"],
    [data-testid="stTextInput"] div[data-baseweb="base-input"] {
        background-color: #FFFFFF !important;
        color: #1F2937 !important;
        -webkit-text-fill-color: #1F2937 !important;
    }
    /* Password reveal (eye) icon */
    [data-testid="stTextInput"] button svg,
    [data-testid="stTextInput"] svg {
        fill: #4B5563 !important;
        color: #4B5563 !important;
    }
    /* "Press Enter to apply" overlay — it overlaps the eye icon and adds no
       value on a single-line field, so hide it there. Kept on textareas,
       where the Cmd+Enter hint is genuinely useful. */
    [data-testid="stTextInput"] [data-testid="InputInstructions"] {
        display: none !important;
    }
    [data-testid="stTextArea"] [data-testid="InputInstructions"] {
        color: #6B7280 !important;
        font-size: 0.68rem !important;
        background: transparent !important;
    }

    /* ── Date input (pre-game game date) ────────────────────────────────── */
    [data-testid="stDateInput"] input,
    [data-testid="stDateInput"] div[data-baseweb="input"] {
        background-color: #FFFFFF !important;
        color: #1F2937 !important;
        -webkit-text-fill-color: #1F2937 !important;
    }

    /* ── Radio options (quiz answers) ───────────────────────────────────── */
    [data-testid="stRadio"] label div[data-testid="stMarkdownContainer"] p {
        color: #1F2937 !important;
        -webkit-text-fill-color: #1F2937 !important;
    }
</style>
""", unsafe_allow_html=True)


# =============================================================================
# ACCESS GATE  (v3.2)
# Shared-password login. Protects your Anthropic API key from public abuse.
#
# Set in .streamlit/secrets.toml (local) or Render env / Streamlit secrets:
#     APP_PASSWORD = "your-crew-password"
#
# If APP_PASSWORD is NOT set, the gate is skipped entirely (convenient for
# local development). NEVER deploy publicly without setting APP_PASSWORD.
# =============================================================================

def get_secret(name: str, default=None):
    """
    Read a configuration value from, in order:

      1. OS environment variables  — Render, Docker, Fly, any normal host
      2. Streamlit secrets.toml    — local dev, Streamlit Community Cloud
      3. Nested [anthropic] table  — legacy secrets.toml layout

    IMPORTANT: st.secrets does NOT read arbitrary environment variables. If no
    secrets.toml file exists anywhere on disk, merely *touching* st.secrets
    raises StreamlitSecretNotFoundError — which is not a KeyError, so a narrow
    `except KeyError` will not catch it and the app crashes on boot. That is
    why every st.secrets access below is wrapped in a broad `except Exception`.
    """
    # 1. Environment variable
    val = os.environ.get(name)
    if val:
        return val

    # 2. Flat key in secrets.toml
    try:
        val = st.secrets[name]
        if val:
            return val
    except Exception:
        pass

    # 3. Legacy nested table, e.g. [anthropic] api_key = "..."
    if name == "ANTHROPIC_API_KEY":
        try:
            val = st.secrets["anthropic"]["api_key"]
            if val:
                return val
        except Exception:
            pass

    return default


def running_on_render() -> bool:
    """Render sets RENDER=true in every service environment."""
    return os.environ.get("RENDER", "").lower() in ("true", "1", "yes")


def check_password() -> bool:
    """
    Render a password prompt and return True only once the correct password
    has been entered. Uses hmac.compare_digest for constant-time comparison
    so the check isn't vulnerable to timing attacks.
    """
    expected = get_secret("APP_PASSWORD")

    if not expected:
        # No password configured.
        if running_on_render():
            # FAIL CLOSED. A public host with no password would expose the
            # owner's API key to the open internet — never allow that silently.
            st.error(
                "🔒 **APP_PASSWORD is not set.**\n\n"
                "This app is running on a public host with no access control, "
                "so it has been locked as a precaution.\n\n"
                "Fix: Render Dashboard → your service → **Environment** → "
                "add `APP_PASSWORD`, then **Save Changes**."
            )
            st.stop()
        # Local development — no password needed, allow through.
        return True

    # Already authenticated this session
    if st.session_state.get("_auth_ok", False):
        return True

    # ── Login screen ─────────────────────────────────────────────────────────
    st.markdown("""
    <div style="text-align:center;padding:3rem 2rem 1rem 2rem;">
        <div style="color:#003087;font-size:3rem;font-weight:900;
                    letter-spacing:-1.5px;">🏈 RefBuddy</div>
        <div style="color:#4B5563;font-size:1.1rem;margin-bottom:0.5rem;">
            Built by a ref, for refs</div>
        <div style="color:#6B7280;font-size:0.9rem;">
            Private tool for MSHSL officials — enter your crew password to continue.</div>
    </div>
    """, unsafe_allow_html=True)

    c1, c2, c3 = st.columns([1, 2, 1])
    with c2:
        pw = st.text_input("Password", type="password",
                           key="_pw_input", label_visibility="collapsed",
                           placeholder="Crew password")
        if st.button("Enter", use_container_width=True, key="_pw_submit"):
            if hmac.compare_digest(str(pw), str(expected)):
                st.session_state["_auth_ok"] = True
                st.rerun()
            else:
                st.error("❌ Incorrect password.")
        st.caption("Need access? Contact your assignor or crew chief.")

    return False


if not check_password():
    st.stop()


# =============================================================================
# SESSION USAGE CAP  (v3.2)
# Second layer of cost protection: caps how many video frames one browser
# session can send to the API. Film analysis is by far the most expensive
# operation (~1,500 tokens per frame), so this is where the money goes.
# =============================================================================

MAX_FRAMES_PER_SESSION = 400   # ~600K tokens ≈ a few dollars, generous for one sitting


def frames_budget_left() -> int:
    """How many analysis frames this session has left."""
    used = st.session_state.get("frames_used", 0)
    return max(0, MAX_FRAMES_PER_SESSION - used)


def spend_frames(n: int) -> bool:
    """
    Charge n frames against the session budget.
    Returns False (and shows an error) if the budget is exhausted.
    """
    if frames_budget_left() < n:
        st.error(
            f"⚠️ **Session analysis limit reached** "
            f"({MAX_FRAMES_PER_SESSION} frames). "
            "Refresh the page to start a new session, or reduce your frame range. "
            "This limit keeps API costs predictable."
        )
        return False
    st.session_state["frames_used"] = st.session_state.get("frames_used", 0) + n
    return True


# =============================================================================
# SESSION STATE
# =============================================================================

def _s(k, v):
    if k not in st.session_state:
        st.session_state[k] = v

_s("frames_used", 0)
_s("messages", [])
_s("uploaded_files_content", [])

# Film
_s("fg_frames", [])
_s("fg_labels", [])
_s("fg_source", "")
_s("fg_is_video", False)
_s("fg_fps", 1.0)
_s("fg_result", "")
_s("fg_q", "")
_s("film_frames", [])
_s("film_frame_count", 0)
_s("film_video_name", "")
_s("film_fps_used", 1.0)
_s("film_analysis_result", "")

# RefGrade
_s("rg_frames", [])
_s("rg_frame_count", 0)
_s("rg_video_name", "")
_s("rg_fps_used", 1.0)
_s("rg_result", "")
_s("rg_saved_logs", [])

# Assignor Hub v2.5 — three sub-tabs
_s("ah_sub", "pregame")       # "pregame" | "eval"
_s("ah_eval_result", "")
_s("ah_eval_scope", "")
_s("ah_crew_frames", [])
_s("ah_crew_frame_count", 0)
_s("ah_crew_video_name", "")
_s("ah_crew_result", "")
_s("ah_ref_frames", [])
_s("ah_ref_frame_count", 0)
_s("ah_ref_video_name", "")
_s("ah_ref_result", "")
_s("ah_pregame_result", "")
_s("ah_pregame_logs", [])

# Quiz
_s("quiz_mode", None)
_s("quiz_topic", "Mixed")
_s("quiz_current_q", None)
_s("quiz_answered", False)
_s("quiz_user_answer", None)
_s("quiz_total", 0)
_s("quiz_correct", 0)
_s("quiz_session_topics", [])    # v2.5: track used topics to avoid repeats
_s("tenq_questions", [])
_s("tenq_index", 0)
_s("tenq_answers", [])
_s("tenq_finished", False)
_s("tenq_answered_this", False)
_s("tenq_user_answer", None)
_s("quiz_log", [])


# =============================================================================
# HELPERS — Core
# =============================================================================

def b64(data: bytes) -> str:
    return base64.standard_b64encode(data).decode("utf-8")

# v3.0: model hard-coded to Sonnet — fast, cost-efficient, no user selector needed
MODEL = "claude-sonnet-4-6"

def make_client():
    """
    Create an Anthropic client.

    The key is resolved by get_secret(), which checks the OS environment first
    (Render / Docker / any host) and then secrets.toml (local dev).
    """
    key = get_secret("ANTHROPIC_API_KEY")
    if not key:
        st.error(
            "❌ **ANTHROPIC_API_KEY not found in secrets.**\n\n"
            "**On Render:** Dashboard → your service → **Environment** → "
            "add `ANTHROPIC_API_KEY` → **Save Changes**.\n\n"
            "**Locally:** add it to `.streamlit/secrets.toml`:\n```\n"
            'ANTHROPIC_API_KEY = "sk-ant-..."\n```'
        )
        st.stop()
    return anthropic.Anthropic(api_key=key)

def api_key_ok() -> bool:
    """v3.0: key comes from secrets — always considered OK at runtime.
    If the secret is missing, make_client() will st.stop() with a clear error."""
    return True

def handle_api_error(e: Exception) -> str:
    if isinstance(e, anthropic.AuthenticationError):
        return "❌ Authentication failed. Check ANTHROPIC_API_KEY in Streamlit secrets."
    if isinstance(e, anthropic.RateLimitError):
        return "⚠️ Rate limit reached. Wait a moment and try again."
    if isinstance(e, anthropic.APIConnectionError):
        return "❌ Connection error. Check your internet connection."
    if isinstance(e, anthropic.BadRequestError):
        return (f"❌ Request too large or malformed: {e}\n\n"
                "Try reducing the frame range or using 0.5 fps extraction.")
    return f"❌ Unexpected error: {e}"

def prepare_file_content(uf):
    data = uf.read()
    name = uf.name.lower()
    if name.endswith(".pdf"):
        return {"type": "document",
                "source": {"type": "base64", "media_type": "application/pdf", "data": b64(data)},
                "title": uf.name}
    if name.endswith((".jpg", ".jpeg")):
        return {"type": "image",
                "source": {"type": "base64", "media_type": "image/jpeg", "data": b64(data)}}
    if name.endswith(".png"):
        return {"type": "image",
                "source": {"type": "base64", "media_type": "image/png", "data": b64(data)}}
    if name.endswith(".txt"):
        return {"type": "text",
                "text": f"[File: {uf.name}]\n\n{data.decode('utf-8', errors='replace')}"}
    return None

def stream_chat(client, messages, files, system=None):
    sys_p = system or SYSTEM_PROMPT
    api_msgs = []
    for i, m in enumerate(messages):
        if m["role"] == "user" and i == len(messages) - 1 and files:
            blocks = list(files) + [{"type": "text", "text": m["content"]}]
            api_msgs.append({"role": "user", "content": blocks})
        else:
            api_msgs.append({"role": m["role"], "content": m["content"]})
    with client.messages.stream(
        model=MODEL, max_tokens=4096,
        system=system_blocks(sys_p), messages=api_msgs, temperature=0,
    ) as s:
        yield from s.text_stream

def call_api_sync(prompt: str, system: str, max_tokens: int = 3000) -> str:
    client = make_client()
    resp = client.messages.create(
        model=MODEL, max_tokens=max_tokens,
        system=system_blocks(system),
        messages=[{"role": "user", "content": prompt}], temperature=0,
    )
    return resp.content[0].text

def chat_log_json() -> str:
    return json.dumps({
        "exported_at": datetime.datetime.now().isoformat(),
        "model": MODEL,
        "messages": [{"role": m["role"], "content": m["content"],
                       "timestamp": m.get("timestamp", "")}
                     for m in st.session_state.messages],
    }, indent=2, ensure_ascii=False)


# =============================================================================
# HELPERS — Frame extraction
# =============================================================================

def _asset_data_uri(filename: str) -> str | None:
    """
    Load a bundled image (e.g. Claude.png) from the app directory and return it
    as a base64 data URI for inline <img> use. Returns None if the file is
    missing so the caller can fall back to text — a missing logo should never
    break the page.
    """
    try:
        here = os.path.dirname(os.path.abspath(__file__))
        path = os.path.join(here, filename)
        with open(path, "rb") as f:
            return "data:image/png;base64," + base64.standard_b64encode(f.read()).decode()
    except Exception:
        return None


def image_to_frame_b64(uploaded_image) -> str | None:
    """
    Convert an uploaded still image (JPG/PNG/HEIC-as-JPG) into the same
    base64-JPEG format extract_frames() produces, so photos and video frames
    flow through one identical downstream path.

    Resized to max 1280px wide to control token cost, same as video frames.
    """
    try:
        data = uploaded_image.read()
        arr = np.frombuffer(data, np.uint8)
        img = cv2.imdecode(arr, cv2.IMREAD_COLOR)
        if img is None:
            return None
        h, w = img.shape[:2]
        if w > 1280:
            img = cv2.resize(img, (1280, int(h * 1280 / w)), interpolation=cv2.INTER_AREA)
        ok, buf = cv2.imencode(".jpg", img, [cv2.IMWRITE_JPEG_QUALITY, 88])
        return base64.standard_b64encode(buf).decode("utf-8") if ok else None
    except Exception:
        return None


def extract_frames(video_path: str, fps: float = 1.0) -> list:
    """Extract frames from video at specified fps. Returns list of base64 JPEG strings."""
    if not OPENCV_AVAILABLE:
        raise RuntimeError("opencv-python-headless not available.")
    cap = cv2.VideoCapture(video_path)
    if not cap.isOpened():
        raise RuntimeError(f"Cannot open video: {video_path}")
    native_fps = cap.get(cv2.CAP_PROP_FPS) or 30.0
    interval = max(1, int(round(native_fps / fps)))
    frames, idx = [], 0
    while True:
        ret, frame = cap.read()
        if not ret:
            break
        if idx % interval == 0:
            h, w = frame.shape[:2]
            if w > 1280:
                frame = cv2.resize(frame, (1280, int(h * 1280 / w)), interpolation=cv2.INTER_AREA)
            ok, buf = cv2.imencode(".jpg", frame, [cv2.IMWRITE_JPEG_QUALITY, 85])
            if ok:
                frames.append(base64.standard_b64encode(buf).decode("utf-8"))
        idx += 1
    cap.release()
    return frames

def extract_video_uploaded(uploaded_video, fps: float = 1.0) -> tuple:
    """Write uploaded video bytes to tmp file, extract frames, return (frames, name)."""
    suffix = ".mp4" if uploaded_video.name.lower().endswith(".mp4") else ".mov"
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(uploaded_video.read())
        tmp_path = tmp.name
    frames = extract_frames(tmp_path, fps=fps)
    os.unlink(tmp_path)
    return frames, uploaded_video.name

def build_vision_content(frames_b64, start_idx, end_idx, user_question,
                          video_name, fps_used, preamble_extra="") -> list:
    selected = frames_b64[start_idx: end_idx + 1]
    spf = 1.0 / fps_used
    content = [{"type": "text", "text": (
        f"Game film: {video_name}\n"
        f"Frames: {len(selected)} ({start_idx+1}–{end_idx+1} of {len(frames_b64)}) "
        f"at {fps_used} fps ({spf:.1f}s/frame).\n"
        f"Frame numbering is 1-based. Use 'Frame N' format.\n{preamble_extra}\n"
    )}]
    for i, fb in enumerate(selected):
        fn = start_idx + i + 1
        content.append({"type": "text", "text": f"--- Frame {fn} (~{(fn-1)/fps_used:.1f}s) ---"})
        content.append({"type": "image", "source": {"type": "base64",
                         "media_type": "image/jpeg", "data": fb}})
    content.append({"type": "text", "text": f"\nQuestion:\n{user_question}"})
    return content

def stream_vision(client, content_blocks, system):
    with client.messages.stream(
        model=MODEL, max_tokens=4096, system=system_blocks(system),
        messages=[{"role": "user", "content": content_blocks}], temperature=0,
    ) as s:
        yield from s.text_stream


# =============================================================================
# HELPERS — Quiz engine (v2.5: improved variety, 50/50 TF/MC)
# =============================================================================

def _strip_json_fences(raw: str) -> str:
    raw = raw.strip()
    if raw.startswith("```"):
        parts = raw.split("```")
        raw = parts[1]
        if raw.startswith("json"):
            raw = raw[4:]
    return raw.strip()

def generate_single_question(topic: str, used_topics: list = None) -> dict | None:
    """
    Generate one quiz question with 50/50 MC/TF mix and no recent repeats.
    used_topics: list of recently used topic+question snippets to avoid repetition.
    """
    if not api_key_ok():
        return None

    # Build anti-repeat instruction
    avoid_str = ""
    if used_topics and len(used_topics) > 0:
        recent = used_topics[-5:]  # last 5 topics
        avoid_str = (f"IMPORTANT: Do NOT generate a question about any of these topics "
                     f"that were just asked: {', '.join(recent)}. "
                     "Pick a completely different rule, mechanic, or scenario.\n")

    # Randomly vary TF vs MC — pass a hint in the prompt
    import random
    q_type_hint = "true_false" if random.random() < 0.5 else "multiple_choice"

    topic_str = "" if topic == "Mixed" else f"Topic focus: {topic}. "
    prompt = (
        f"{avoid_str}"
        f"{topic_str}"
        f"Generate one {q_type_hint} question for a Minnesota high school football referee. "
        f"It must be challenging, specific, and reference exact rule numbers. "
        f"For multiple_choice: EXACTLY 4 options (A, B, C, D). "
        f"For true_false: EXACTLY 2 options (A=True, B=False). "
        f"Respond with ONLY valid JSON — no fences, no preamble."
    )
    try:
        client = make_client()
        resp = client.messages.create(
            model=MODEL, max_tokens=900,
            system=system_blocks(QUIZ_SYSTEM_PROMPT),
            messages=[{"role": "user", "content": prompt}],
            temperature=0,
        )
        raw = _strip_json_fences(resp.content[0].text)
        q = json.loads(raw)
        # Validate
        if q.get("type") == "multiple_choice" and len(q.get("options", {})) != 4:
            return None  # wrong count — caller will retry
        if q.get("type") == "true_false" and len(q.get("options", {})) != 2:
            return None
        return q
    except Exception as e:
        st.error(f"❌ Failed to generate question: {e}")
        return None

def generate_ten_questions(topic: str) -> list | None:
    """Generate 10 varied quiz questions with 50/50 MC/TF mix."""
    if not api_key_ok():
        return None
    topic_str = "" if topic == "Mixed" else f"Topic focus: {topic}. "
    prompt = (
        f"{topic_str}Generate exactly 10 questions for a Minnesota high school football referee. "
        "Mix: 5 multiple_choice (EXACTLY 4 options A/B/C/D each) + 5 true_false. "
        "Cover these areas: 2026 rule changes, MSHSL mercy rule, BJ mechanics, "
        "penalty enforcement, player eligibility, kick rules, personal game note scenarios, "
        "USC/sportsmanship, timing rules, and safety/free kick. "
        "Respond with ONLY a valid JSON array of 10 objects — no fences, no preamble."
    )
    try:
        client = make_client()
        resp = client.messages.create(
            model=MODEL, max_tokens=6000,
            system=system_blocks(QUIZ_SYSTEM_PROMPT),
            messages=[{"role": "user", "content": prompt}],
            temperature=0,
        )
        raw = _strip_json_fences(resp.content[0].text)
        questions = json.loads(raw)
        if isinstance(questions, list) and len(questions) == 10:
            return questions
        # Try to salvage partial list
        if isinstance(questions, list) and len(questions) >= 5:
            return questions[:10]
        st.error("❌ Unexpected question count. Try again.")
        return None
    except Exception as e:
        st.error(f"❌ Failed to generate quiz: {e}")
        return None

def render_question_card(q: dict, question_num: str = ""):
    q_text = q.get("question", "")
    q_type = q.get("type", "multiple_choice")
    badge_color = "#EEF2FF"
    badge_label = "True/False" if q_type == "true_false" else "Multiple Choice"
    st.markdown(f"""
    <div class="quiz-question-card">
        <div class="quiz-question-text">{question_num} {q_text}
        <span style="background:{badge_color};color:{BLUE};font-size:0.72rem;
        font-weight:700;border-radius:20px;padding:2px 8px;margin-left:8px;">
        {badge_label}</span></div>
    </div>
    """, unsafe_allow_html=True)

def render_feedback(q: dict, user_answer: str) -> bool:
    correct = q.get("correct", "")
    options = q.get("options", {})
    correct_text = options.get(correct, correct)
    user_text = options.get(user_answer, user_answer)
    is_correct = user_answer == correct

    if is_correct:
        st.markdown(f"""<div class="quiz-result-correct">
        <strong>✅ Correct!</strong> &nbsp; {user_answer}: {user_text}
        </div>""", unsafe_allow_html=True)
    else:
        st.markdown(f"""<div class="quiz-result-wrong">
        <strong>❌ Incorrect.</strong> You chose: {user_answer}: {user_text}<br>
        <strong>✔ Correct: {correct}: {correct_text}</strong>
        </div>""", unsafe_allow_html=True)

    explanation = q.get("explanation", "")
    rule_cite = q.get("rule_citation", "")
    personal = q.get("personal_note", "")
    pnote = f'<br><strong>📋 From your notes:</strong> {personal}' if personal else ""
    st.markdown(f"""<div class="quiz-explanation">
    <strong>📖 Explanation</strong><br>{explanation}<br><br>
    <strong>📌 Citation:</strong> {rule_cite}{pnote}
    </div>""", unsafe_allow_html=True)
    return is_correct

def accuracy_display(correct: int, total: int):
    pct = int(round(correct / total * 100)) if total > 0 else 0
    color = "#15803D" if pct >= 80 else ("#92400E" if pct >= 60 else "#991B1B")
    st.markdown(f"""
    <div style="display:flex;align-items:center;gap:12px;background:{CARD};
                border:1px solid {BORDER};border-radius:8px;padding:0.7rem 1rem;
                margin-bottom:0.8rem;">
        <div style="font-weight:800;font-size:1.4rem;color:{color};min-width:52px;">{pct}%</div>
        <div style="flex:1;">
            <div class="accuracy-bar-wrap">
                <div class="accuracy-bar-fill" style="width:{pct}%;background:{color};"></div>
            </div>
            <div style="font-size:0.8rem;color:{MUTED};margin-top:3px;">
                {correct} correct of {total} answered</div>
        </div>
    </div>""", unsafe_allow_html=True)


# =============================================================================
# HELPERS — Export (PDF + DOCX)
# =============================================================================

def sanitize_for_pdf(text: str) -> str:
    """
    Replace non-latin-1 characters that cause FPDFUnicodeEncodingException.
    fpdf2 with core fonts (Helvetica) is limited to latin-1 / CP1252.
    This maps the most common Unicode punctuation and symbols to safe ASCII.
    Applied to EVERY string before any pdf.multi_cell() or pdf.cell() call.
    """
    if not text:
        return ""
    replacements = {
        # Dashes and hyphens
        "\u2014": "-",   # em dash —
        "\u2013": "-",   # en dash –
        "\u2012": "-",   # figure dash
        "\u2015": "-",   # horizontal bar
        # Quotes
        "\u2018": "'",   # left single '
        "\u2019": "'",   # right single ' / apostrophe
        "\u201a": ",",   # single low-9 quotation ,
        "\u201b": "'",   # single high-reversed-9
        "\u201c": '"',   # left double "
        "\u201d": '"',   # right double "
        "\u201e": '"',   # double low-9 quotation
        "\u201f": '"',   # double high-reversed-9
        # Ellipsis and bullets
        "\u2026": "...", # horizontal ellipsis …
        "\u2022": "*",   # bullet •
        "\u2023": ">",   # triangular bullet
        "\u25e6": "o",   # white bullet
        "\u2043": "-",   # hyphen bullet
        # Spaces and misc
        "\u00a0": " ",   # non-breaking space
        "\u200b": "",    # zero-width space
        "\u200c": "",    # zero-width non-joiner
        "\u200d": "",    # zero-width joiner
        "\u2060": "",    # word joiner
        "\ufeff": "",    # BOM
        # Math / technical
        "\u00d7": "x",   # multiplication sign ×
        "\u00f7": "/",   # division sign ÷
        "\u2212": "-",   # minus sign −
        "\u00b0": "°",   # degree sign (latin-1 safe)
        # Arrows (common in AI output)
        "\u2192": "->",  # rightwards arrow →
        "\u2190": "<-",  # leftwards arrow ←
        "\u2194": "<->", # left-right arrow ↔
        # Emoji — strip them entirely
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    # Final pass: strip any remaining non-latin-1 characters
    text = text.encode("latin-1", errors="replace").decode("latin-1")
    return text.strip()


def markdown_to_pdf_bytes(md_text: str, title: str = "RefBuddy Report") -> bytes | None:
    """
    Convert markdown to PDF via fpdf2.
    v2.7 improvements over v2.6:
      - Wider margins (20mm L/R) so text never clips
      - Smaller body font (9pt) with tighter line height (4.5mm) for compact output
      - Explicit effective_width passed to every multi_cell so text wraps correctly
        instead of overflowing on some fpdf2 versions
      - sanitize_for_pdf() on every string (unchanged from v2.6)
    Returns bytes or None if fpdf2 is not installed.
    """
    try:
        from fpdf import FPDF  # type: ignore

        L_MARGIN = 20   # mm — wider than default (10) to prevent clipping
        R_MARGIN = 20
        TOP_MARGIN = 18
        BOT_MARGIN = 15

        pdf = FPDF()
        pdf.set_margins(L_MARGIN, TOP_MARGIN, R_MARGIN)
        pdf.set_auto_page_break(auto=True, margin=BOT_MARGIN)
        pdf.add_page()

        # Effective text width between margins
        eff_w = pdf.w - L_MARGIN - R_MARGIN  # ~170mm on A4

        # ── Title ─────────────────────────────────────────────────────────────
        pdf.set_font("Helvetica", "B", 14)
        pdf.set_text_color(0, 48, 135)
        pdf.multi_cell(eff_w, 7, sanitize_for_pdf(title), align="L")
        pdf.ln(1)

        # ── Date line ─────────────────────────────────────────────────────────
        pdf.set_font("Helvetica", "", 8)
        pdf.set_text_color(107, 114, 128)
        date_str = sanitize_for_pdf(
            f"Generated: {datetime.datetime.now().strftime('%B %d, %Y %H:%M')}"
        )
        pdf.multi_cell(eff_w, 4, date_str, align="L")
        pdf.ln(2)

        # ── Divider ───────────────────────────────────────────────────────────
        pdf.set_draw_color(180, 200, 220)
        pdf.set_line_width(0.3)
        pdf.line(L_MARGIN, pdf.get_y(), pdf.w - R_MARGIN, pdf.get_y())
        pdf.ln(3)

        # ── Body — reset to 9pt dark text ─────────────────────────────────────
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(31, 41, 55)

        for raw_line in md_text.split("\n"):
            s = raw_line.strip()

            # Blank line → small vertical gap (not a full line)
            if not s:
                pdf.ln(2)
                continue

            # Horizontal rule
            if s == "---":
                pdf.ln(1)
                pdf.set_draw_color(200, 210, 220)
                pdf.line(L_MARGIN, pdf.get_y(), pdf.w - R_MARGIN, pdf.get_y())
                pdf.ln(2)
                continue

            # H1 / H2
            if s.startswith("# ") or s.startswith("## "):
                pdf.ln(2)
                pdf.set_font("Helvetica", "B", 11)
                pdf.set_text_color(0, 48, 135)
                pdf.multi_cell(eff_w, 6, sanitize_for_pdf(s.lstrip("#").strip()), align="L")
                pdf.set_font("Helvetica", "", 9)
                pdf.set_text_color(31, 41, 55)
                continue

            # H3
            if s.startswith("### "):
                pdf.ln(1)
                pdf.set_font("Helvetica", "B", 10)
                pdf.set_text_color(0, 48, 135)
                pdf.multi_cell(eff_w, 5, sanitize_for_pdf(s.lstrip("#").strip()), align="L")
                pdf.set_font("Helvetica", "", 9)
                pdf.set_text_color(31, 41, 55)
                continue

            # H4
            if s.startswith("#### "):
                pdf.ln(1)
                pdf.set_font("Helvetica", "B", 9)
                pdf.set_text_color(31, 41, 55)
                pdf.multi_cell(eff_w, 5, sanitize_for_pdf(s.lstrip("#").strip()), align="L")
                pdf.set_font("Helvetica", "", 9)
                continue

            # Bullet / list item — indent 6mm, hanging
            if s.startswith(("- ", "* ", "+ ")):
                content = sanitize_for_pdf(s[2:].replace("**", ""))
                x_orig = pdf.get_x()
                # Print bullet marker at left margin
                pdf.set_x(L_MARGIN)
                pdf.cell(6, 4.5, "-")
                # Print content with remaining width
                pdf.set_x(L_MARGIN + 6)
                pdf.multi_cell(eff_w - 6, 4.5, content, align="L")
                continue

            # Numbered list item (e.g. "1. text")
            import re as _re
            num_match = _re.match(r"^(\d+)\.\s+(.*)", s)
            if num_match:
                num = num_match.group(1) + "."
                content = sanitize_for_pdf(num_match.group(2).replace("**", ""))
                pdf.set_x(L_MARGIN)
                pdf.cell(8, 4.5, num)
                pdf.set_x(L_MARGIN + 8)
                pdf.multi_cell(eff_w - 8, 4.5, content, align="L")
                continue

            # Table row — strip pipes, use Courier 8pt
            if s.startswith("| ") or (s.startswith("|") and "|" in s[1:]):
                # Skip pure separator rows like |---|---|
                stripped = s.replace("|", "").replace("-", "").replace(" ", "")
                if not stripped:
                    continue
                row = sanitize_for_pdf(
                    s.strip("|").replace("|", "  ").replace("**", "")
                )
                pdf.set_font("Courier", "", 8)
                pdf.multi_cell(eff_w, 4, row, align="L")
                pdf.set_font("Helvetica", "", 9)
                continue

            # Plain body text
            clean = sanitize_for_pdf(s.replace("**", "").replace("*", ""))
            pdf.multi_cell(eff_w, 4.5, clean, align="L")

        # fpdf2's pdf.output() returns bytearray on some versions and bytes on others.
        # bytes() normalises both to a plain bytes object that st.download_button accepts.
        return bytes(pdf.output())

    except ImportError:
        return None

def markdown_to_docx_bytes(md_text: str, title: str = "RefBuddy Report") -> bytes | None:
    """
    Convert markdown text to a Word .docx file via python-docx.
    Returns bytes or None if python-docx is not installed.
    """
    try:
        from docx import Document  # type: ignore
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Page margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.2)
            section.right_margin = Inches(1.2)

        # Title
        t_para = doc.add_paragraph()
        t_run = t_para.add_run(title)
        t_run.bold = True
        t_run.font.size = Pt(18)
        t_run.font.color.rgb = RGBColor(0, 48, 135)
        t_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Date
        d_para = doc.add_paragraph()
        d_run = d_para.add_run(
            f"Generated: {datetime.datetime.now().strftime('%B %d, %Y %H:%M')}"
        )
        d_run.font.size = Pt(9)
        d_run.font.color.rgb = RGBColor(107, 114, 128)
        doc.add_paragraph()

        # Body
        for line in md_text.split("\n"):
            s = line.strip()
            if s.startswith("## ") or s.startswith("# "):
                h = doc.add_heading(s.lstrip("#").strip(), level=2)
                for run in h.runs:
                    run.font.color.rgb = RGBColor(0, 48, 135)
            elif s.startswith("### "):
                h = doc.add_heading(s.lstrip("#").strip(), level=3)
                for run in h.runs:
                    run.font.color.rgb = RGBColor(0, 48, 135)
            elif s.startswith(("- ", "* ")):
                content = s[2:].replace("**", "")
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(content).font.size = Pt(10)
            elif s == "---":
                doc.add_paragraph()
                doc.add_paragraph().add_run("─" * 60).font.size = Pt(8)
                doc.add_paragraph()
            elif s == "":
                doc.add_paragraph()
            else:
                clean = s.replace("**", "")
                p = doc.add_paragraph()
                p.add_run(clean).font.size = Pt(10)

        # Save to bytes
        buf = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        buf.close()
        doc.save(buf.name)
        with open(buf.name, "rb") as f:
            data = f.read()
        os.unlink(buf.name)
        return data
    except ImportError:
        return None


# =============================================================================
# SIDEBAR
# =============================================================================

with st.sidebar:
    # v3.0: light bg brand block
    st.markdown(
        '<div style="background:#F8FAFC;border:2px solid #1F2937;border-radius:8px;'
        'padding:0.7rem 1rem;margin-bottom:0.8rem;">'
        '<span style="color:#1F2937;font-weight:800;font-size:1.1rem;">🏈 RefBuddy</span><br>'
        '<span style="color:#4B5563;font-size:0.72rem;">Built by a ref, for refs</span>'
        '</div>',
        unsafe_allow_html=True,
    )

    # v3.5: "Powered by Claude" lockup (matches PickleRick styling)
    _claude_uri = _asset_data_uri("Claude.png")
    if _claude_uri:
        st.markdown(
            '<div style="display:flex;align-items:center;gap:12px;'
            'flex-wrap:wrap;margin:0.7rem 0 0.4rem 0;">'
            '<span style="color:#1F2937;font-weight:700;font-size:1.05rem;'
            'line-height:1;">Powered by</span>'
            f'<img src="{_claude_uri}" alt="Claude" '
            'style="height:85px;width:auto;display:block;">'
            '</div>',
            unsafe_allow_html=True,
        )
    else:
        st.markdown(
            '<div style="color:#1F2937;font-weight:700;font-size:0.9rem;'
            'margin:0.5rem 0 0.2rem 0;">Powered by Claude</div>',
            unsafe_allow_html=True,
        )

    # v3.2: session analysis budget indicator
    _left = frames_budget_left()
    _pct = int(_left / MAX_FRAMES_PER_SESSION * 100)
    _cls = "pill-ok" if _pct > 40 else ("pill-warn" if _pct > 15 else "pill-err")
    st.markdown(
        f'<span class="{_cls}">🎬 {_left} analysis frames left</span>',
        unsafe_allow_html=True,
    )

    st.markdown("---")
    st.markdown("**Knowledge Base**")
    st.caption("Years of NFHS veteran officials' game notes and NFHS/MSHSL rulebook facts and interpretations")

    st.markdown("---")
    st.markdown("**Upload Files** *(Home chat)*")
    st.caption("PDFs, images, or TXT")
    chat_uploads = st.file_uploader(
        "chatfiles", type=["pdf", "jpg", "jpeg", "png", "txt"],
        accept_multiple_files=True, label_visibility="collapsed",
    )
    if chat_uploads:
        proc, names = [], []
        for uf in chat_uploads:
            c = prepare_file_content(uf)
            if c:
                proc.append(c); names.append(uf.name)
            else:
                st.warning(f"Unsupported: {uf.name}")
        st.session_state.uploaded_files_content = proc
        if names:
            st.markdown(f'<span class="pill-ok">✅ {len(names)} file(s)</span>', unsafe_allow_html=True)
    else:
        st.session_state.uploaded_files_content = []

    st.markdown("---")
    st.markdown("**Ref Log**")
    if st.session_state.messages:
        ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        st.download_button("⬇️ Download Chat Log", data=chat_log_json(),
                           file_name=f"refbuddy_chat_{ts}.json",
                           mime="application/json", use_container_width=True)
    if st.button("🗑️ Clear Home Chat", use_container_width=True):
        st.session_state.messages = []
        st.rerun()

# =============================================================================
# TABS
# =============================================================================

tab_home, tab_film, tab_ah, tab_quiz = st.tabs([
    "🏈 Home",
    "🎬 Film & Grade",
    "👥 Ref Hub",
    "📝 Quiz",
])


# ─────────────────────────────────────────────────────────────────────────────
# TAB 1 — HOME
# ─────────────────────────────────────────────────────────────────────────────

with tab_home:
    st.markdown("""
    <div class="home-hero">
        <div class="home-hero-title">🏈 RefBuddy</div>
        <div class="home-hero-slogan">Built by a ref, for refs</div>
    </div>
    """, unsafe_allow_html=True)

    chips = ["NFHS Rule Citations", "MSHSL Mods", "Penalty Enforcement",
             "BJ · U · Wings · R", "Mercy Rule", "Film Analysis", "Quiz"]
    chip_html = " &nbsp; ".join(f'<span class="pill-blue">{c}</span>' for c in chips)
    st.markdown(f'<div style="text-align:center;margin-bottom:1.4rem;line-height:2.6;">'
                f'{chip_html}</div>', unsafe_allow_html=True)

    # ── Quick-start prompts (only shown when chat is empty) ───────────────────
    if not st.session_state.messages:
        st.markdown(f'<p style="text-align:center;color:{MUTED};font-size:0.9rem;'
                    f'margin-bottom:0.8rem;"><em>Try one of these or type your own below</em></p>',
                    unsafe_allow_html=True)
        starter_qs = [
            "What's the penalty for 12 players on the field in NFHS vs NCAA?",
            "Walk me through Back Judge responsibilities on a punt.",
            "When does the Minnesota mercy rule clock stop?",
            "What's the correct call when a forward fumble goes out of bounds?",
            "Explain targeting vs. spearing vs. defenseless receiver contact.",
            "What are my kickoff positioning duties as Back Judge?",
            "When does the play clock reset to 25 vs 40 seconds?",
            "What are the 2026 NFHS rule changes I need to know?",
        ]
        c1, c2 = st.columns(2)
        for i, q in enumerate(starter_qs):
            col = c1 if i < 4 else c2
            with col:
                if st.button(f"➤ {q}", key=f"hq_{i}", use_container_width=True):
                    st.session_state.messages.append({
                        "role": "user", "content": q,
                        "timestamp": datetime.datetime.now().isoformat(),
                    })
                    st.rerun()

    # ── Render full chat history ──────────────────────────────────────────────
    # All messages (user + assistant) render here in a scrollable column.
    # st.chat_input() is called LAST so Streamlit pins it to the bottom of the
    # viewport — nothing ever renders below it.
    for msg in st.session_state.messages:
        with st.chat_message(msg["role"], avatar="🏈" if msg["role"] == "user" else "⚡"):
            st.markdown(msg["content"])

    # ── Stream assistant reply for the last unanswered user message ───────────
    # This block runs BEFORE st.chat_input() so the response streams into the
    # scrollable history area above the fixed input box, never below it.
    if (st.session_state.messages
            and st.session_state.messages[-1]["role"] == "user"):
        pending = st.session_state.messages[-1]["content"]
        if not api_key_ok():
            st.warning("⚠️ Enter your Anthropic API key in the sidebar.")
        else:
            client = make_client()
            with st.chat_message("assistant", avatar="⚡"):
                ph = st.empty()
                full = ""
                try:
                    for chunk in stream_chat(
                        client,
                        st.session_state.messages,
                        st.session_state.uploaded_files_content,
                    ):
                        full += chunk
                        ph.markdown(full + "▌")
                    ph.markdown(full)
                    st.session_state.messages.append({
                        "role": "assistant", "content": full,
                        "timestamp": datetime.datetime.now().isoformat(),
                    })
                except Exception as e:
                    st.error(handle_api_error(e))

    # ── Single chat input — Streamlit pins this to the bottom of the viewport ─
    user_in = st.chat_input(
        "Ask anything about NFHS/MSHSL rules, game situations, or your notes…",
    )
    if user_in:
        st.session_state.messages.append({
            "role": "user", "content": user_in,
            "timestamp": datetime.datetime.now().isoformat(),
        })
        st.rerun()  # triggers the streaming block above on next render

    # ── Ref Log expander ──────────────────────────────────────────────────────
    if st.session_state.messages:
        st.markdown("---")
        with st.expander("📋 Ref Log — Session Summary", expanded=False):
            st.markdown(f"""<div class="ref-log">
            <strong>Session Stats</strong><br>
            Messages: {len(st.session_state.messages)} &nbsp;|&nbsp;
            Model: {MODEL}<br>
            Started: {st.session_state.messages[0].get("timestamp","")[:19]}<br>
            Last: {st.session_state.messages[-1].get("timestamp","")[:19]}
            </div>""", unsafe_allow_html=True)
            for i, m in enumerate(st.session_state.messages):
                icon = "🏈 You" if m["role"] == "user" else "⚡ RefBuddy"
                st.markdown(f"**{icon}** _{m.get('timestamp','')[:19]}_")
                st.markdown(m["content"])
                if i < len(st.session_state.messages) - 1:
                    st.markdown("---")
            ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            st.download_button("⬇️ Save Ref Log", data=chat_log_json(),
                               file_name=f"refbuddy_reflog_{ts}.json",
                               mime="application/json")


# ─────────────────────────────────────────────────────────────────────────────
# TAB 2 — GAME FILM ANALYZER
# ─────────────────────────────────────────────────────────────────────────────
# TAB 2 — FILM & GRADE  (v3.5: Game Film + RefGrade merged into one tab)
#
# Accepts BOTH still images (a photo of a call, a screenshot from Hudl) and
# video clips. Photos are the common case — they're easy to capture on a phone
# and easy to upload. Video is still supported for crews who have clips handy.
#
# Two analysis modes share the same uploaded media:
#   • Ask a Question  — free-form rules/mechanics analysis (SYSTEM_PROMPT)
#   • RefGrade        — structured scored evaluation (REFGRADE_PROMPT)
# ─────────────────────────────────────────────────────────────────────────────

with tab_film:
    st.markdown("## 🎬 Film & Grade")
    st.markdown(
        "Upload **photos of a call** or a **short video clip**, then either ask a "
        "question about it or run a full scored RefGrade evaluation. "
        "Every analysis begins with a Visibility Check so you know what was "
        "actually observed versus inferred."
    )

    # ── Step 1 — Upload ──────────────────────────────────────────────────────
    st.markdown("### Step 1 — Upload")
    st.info(
        "📸 **Photos work best** — screenshots from film, a still of the formation, "
        "or a photo of the play. Upload as many as you like.\n\n"
        "🎥 **Video** (.mp4/.mov) also works. Keep clips to 10–60 seconds and trim "
        "to the play. Large files can be slow or fail to upload on phone data — "
        "if a video won't go through, screenshot the key moments instead."
    )

    fg_uploads = st.file_uploader(
        "fg_upload",
        type=["jpg", "jpeg", "png", "mp4", "mov"],
        accept_multiple_files=True,
        label_visibility="collapsed",
        key="fg_uploader",
    )

    if fg_uploads:
        images = [f for f in fg_uploads
                  if f.name.lower().endswith((".jpg", ".jpeg", ".png"))]
        videos = [f for f in fg_uploads
                  if f.name.lower().endswith((".mp4", ".mov"))]

        # fps control only matters when a video is present
        fg_fps = 1.0
        if videos:
            fg_fps = st.select_slider(
                "Video sampling rate (frames per second)",
                options=[0.5, 1.0, 2.0], value=1.0,
                help="0.5 = overview | 1.0 = standard | 2.0 = fast action",
                key="fg_fps_slider",
            )
            st.caption(f"A 30s clip at {fg_fps} fps ≈ {int(30*fg_fps)} frames")

        if st.button("📥 Load Media for Analysis", use_container_width=True,
                     key="fg_load"):
            frames, labels = [], []
            with st.spinner("Processing upload…"):
                # Stills first — they keep their original order
                for img in images:
                    b64_img = image_to_frame_b64(img)
                    if b64_img:
                        frames.append(b64_img)
                        labels.append(img.name)
                    else:
                        st.warning(f"Could not read image: {img.name}")

                # Then any video frames
                for vid in videos:
                    if not OPENCV_AVAILABLE:
                        st.error("Video support requires opencv-python-headless.")
                        break
                    try:
                        suffix = ".mp4" if vid.name.lower().endswith(".mp4") else ".mov"
                        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                            tmp.write(vid.read()); tmp_path = tmp.name
                        vframes = extract_frames(tmp_path, fps=fg_fps)
                        os.unlink(tmp_path)
                        if not vframes:
                            st.error(f"No frames extracted from {vid.name} — "
                                     "check the file format or codec.")
                        frames.extend(vframes)
                        labels.extend([f"{vid.name} frame {i+1}"
                                       for i in range(len(vframes))])
                    except Exception as e:
                        st.error(f"❌ Could not process {vid.name}: {e}")

            if frames:
                st.session_state.fg_frames = frames
                st.session_state.fg_labels = labels
                st.session_state.fg_source = ", ".join(
                    [f.name for f in fg_uploads][:3]
                ) + ("…" if len(fg_uploads) > 3 else "")
                st.session_state.fg_is_video = bool(videos)
                st.session_state.fg_fps = fg_fps
                st.session_state.fg_result = ""
                st.success(
                    f"✅ Loaded {len(frames)} image(s) — "
                    f"{len(images)} photo(s), {len(frames)-len(images)} video frame(s)"
                )
            else:
                st.error("Nothing could be loaded from that upload.")

    # ── Steps 2-4 — only once media is loaded ────────────────────────────────
    if st.session_state.get("fg_frames"):
        frames = st.session_state.fg_frames
        n = len(frames)
        fps_u = st.session_state.get("fg_fps", 1.0)
        is_vid = st.session_state.get("fg_is_video", False)
        src_name = st.session_state.get("fg_source", "upload")

        st.markdown("---")
        st.markdown(f"**{n} image(s) loaded** from `{src_name}`")

        # Range selector — only useful with many frames
        if n > 1:
            st.markdown("### Step 2 — Select Range")
            sf, ef = st.slider("fg_range", 1, n, (1, min(n, 30)), key="fg_range_sel",
                               label_visibility="collapsed")
        else:
            sf = ef = 1
        sel = ef - sf + 1
        st.caption(f"Analyzing image {sf}–{ef} ({sel} total)")

        with st.expander(f"🔍 Preview {sel} selected", expanded=(sel <= 12)):
            cols = st.columns(4)
            for i, fb in enumerate(frames[sf-1:ef][:24]):
                with cols[i % 4]:
                    st.image(base64.b64decode(fb), caption=f"#{sf+i}",
                             use_container_width=True)

        # ── Step 3 — Choose what to do ───────────────────────────────────────
        st.markdown("### Step 3 — What do you want?")
        fg_mode = st.radio(
            "fg_mode",
            options=["❓ Ask a Question", "📊 RefGrade Evaluation"],
            horizontal=True,
            label_visibility="collapsed",
            key="fg_mode_sel",
        )

        # ══════════════════════════════════════════════════════════════════
        # MODE A — Ask a Question
        # ══════════════════════════════════════════════════════════════════
        if fg_mode == "❓ Ask a Question":
            st.markdown("**Quick presets**")
            p1, p2, p3 = st.columns(3)
            PRESETS = {
                "flag": ("Analyze this for any rule violations under NFHS rules. "
                         "For each potential foul: cite the rule number, describe what "
                         "you see using 'Frame N' format, state the correct penalty and "
                         "enforcement spot, and note which official had primary "
                         "responsibility. Begin with a VISIBILITY CHECK."),
                "mech": ("Evaluate the officiating mechanics visible here. Begin with a "
                         "VISIBILITY CHECK. For each visible official describe their "
                         "positioning, whether it matches the mechanics manual, and any "
                         "improvements. Reference specific frame numbers."),
                "form": ("Analyze the offensive and defensive formations. Check Rule 7-2-5 "
                         "(5 players numbered 50-79 on the LOS, max 4 backs), motion "
                         "legality per Rule 7-2-7, and defensive alignment. "
                         "Begin with a VISIBILITY CHECK."),
            }
            # A preset writes straight into session_state and reruns.
            # Passing value= to a widget that already has a key is a no-op on
            # rerun — Streamlit keeps the stored widget state — which is why
            # these buttons previously appeared to do nothing.
            with p1:
                if st.button("🚩 Was there a foul?", key="fg_p_flag",
                             use_container_width=True):
                    st.session_state["fg_q"] = PRESETS["flag"]
                    st.rerun()
            with p2:
                if st.button("⚙️ Check mechanics", key="fg_p_mech",
                             use_container_width=True):
                    st.session_state["fg_q"] = PRESETS["mech"]
                    st.rerun()
            with p3:
                if st.button("📐 Check formation", key="fg_p_form",
                             use_container_width=True):
                    st.session_state["fg_q"] = PRESETS["form"]
                    st.rerun()

            fg_q = st.text_area(
                "fg_question", height=110,
                placeholder="e.g. 'Is the left tackle holding here?' or "
                            "'Was my positioning correct as Back Judge on this play?'",
                label_visibility="collapsed", key="fg_q",
            )

            if st.button(f"🔍  Analyze {sel} Image{'s' if sel != 1 else ''}",
                         disabled=not (fg_q or "").strip(),
                         type="primary",
                         use_container_width=True, key="fg_run_q"):
                if not spend_frames(sel):
                    st.stop()
                blocks = build_vision_content(
                    frames, sf-1, ef-1, fg_q.strip(), src_name, fps_u,
                    preamble_extra=(
                        "These may be still photos rather than sequential video frames. "
                        "Begin with a VISIBILITY CHECK listing which crew members are "
                        "clearly visible, partially visible, or not visible."
                    ),
                )
                st.markdown("---")
                st.markdown("#### ⚡ Analysis")
                ph = st.empty(); full = ""
                try:
                    with st.spinner(f"Analyzing {sel} image(s)…"):
                        for chunk in stream_vision(make_client(), blocks, SYSTEM_PROMPT):
                            full += chunk; ph.markdown(full + "▌")
                    ph.markdown(full)
                    st.session_state.fg_result = full
                except Exception as e:
                    st.error(handle_api_error(e))

        # ══════════════════════════════════════════════════════════════════
        # MODE B — RefGrade Evaluation
        # ══════════════════════════════════════════════════════════════════
        else:
            st.markdown("**Evaluation setup**")
            g1, g2 = st.columns(2)
            with g1:
                fg_scope = st.selectbox(
                    "Evaluate", options=[
                        "Full Crew (Overall)", "Referee (R)", "Umpire (U)",
                        "Line Judge (LJ)", "Down Judge (DJ)", "Back Judge (BJ)",
                        "Side Judge (SJ)", "Field Judge (FJ)"],
                    key="fg_scope")
            with g2:
                fg_crew = st.selectbox(
                    "Crew size",
                    options=["3-Person Crew", "4-Person Crew", "5-Person Crew"],
                    key="fg_crew")

            fg_cats = st.multiselect(
                "Score these categories", options=[
                    "Positioning", "Call Accuracy", "Mechanics Execution",
                    "Dead-ball Officiating", "Communication / Signals",
                    "Clock Management", "Penalty Administration"],
                default=["Positioning", "Call Accuracy", "Mechanics Execution",
                         "Communication / Signals"],
                key="fg_cats")

            fg_notes = st.text_area(
                "Context or focus (optional)", height=80,
                placeholder="e.g. 'A flag was thrown on this play — was it correct?' "
                            "or 'Check my depth as Back Judge.'",
                key="fg_notes")

            if st.button(f"📊  Run RefGrade — {fg_scope}",
                         disabled=not fg_cats,
                         type="primary",
                         use_container_width=True, key="fg_run_grade"):
                if not spend_frames(sel):
                    st.stop()
                cats = ", ".join(fg_cats)
                extra = f"\nContext: {fg_notes.strip()}" if fg_notes.strip() else ""
                gq = (f"Perform a RefGrade evaluation.\n"
                      f"Source: {src_name}\nScope: {fg_scope}\nCrew: {fg_crew}\n"
                      f"Score these categories: {cats}{extra}\n\n"
                      f"Use the exact RefGrade report structure from your system prompt. "
                      f"Begin with a VISIBILITY CHECK. Cite the specific mechanic or rule "
                      f"behind every score. If these are still photos rather than video, "
                      f"say so and scope your confidence accordingly.")
                blocks = build_vision_content(
                    frames, sf-1, ef-1, gq, src_name, fps_u,
                    preamble_extra=("Structured RefGrade evaluation. Visibility Check is "
                                    "the mandatory first section. These may be stills "
                                    "rather than sequential frames."),
                )
                st.markdown("---")
                st.markdown("#### 📊 RefGrade Report")
                ph = st.empty(); full = ""
                try:
                    with st.spinner("Running RefGrade… (20–90 seconds)"):
                        for chunk in stream_vision(make_client(), blocks, REFGRADE_PROMPT):
                            full += chunk; ph.markdown(full + "▌")
                    ph.markdown(full)
                    st.session_state.fg_result = full
                    st.session_state.rg_saved_logs.append({
                        "timestamp": datetime.datetime.now().isoformat(),
                        "source": src_name, "scope": fg_scope, "crew": fg_crew,
                        "range": f"{sf}-{ef}", "result": full,
                    })
                except Exception as e:
                    st.error(handle_api_error(e))

        # ── Step 4 — Export ──────────────────────────────────────────────────
        if st.session_state.get("fg_result"):
            st.markdown("---")
            e1, e2, e3 = st.columns(3)
            ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            with e1:
                st.download_button("⬇️ Download TXT",
                                   data=st.session_state.fg_result,
                                   file_name=f"refbuddy_analysis_{ts}.txt",
                                   mime="text/plain", use_container_width=True)
            with e2:
                _pdf = markdown_to_pdf_bytes(st.session_state.fg_result,
                                             "RefBuddy Film Analysis")
                if _pdf:
                    st.download_button("⬇️ Export PDF", data=_pdf,
                                       file_name=f"refbuddy_analysis_{ts}.pdf",
                                       mime="application/pdf",
                                       use_container_width=True)
            with e3:
                if st.button("🗑️ Clear", use_container_width=True, key="fg_clear"):
                    for k in ("fg_frames", "fg_labels", "fg_source",
                              "fg_is_video", "fg_result"):
                        st.session_state.pop(k, None)
                    st.rerun()

    elif not fg_uploads:
        st.markdown("---")
        st.markdown("""<div class="rb-card-blue">
        <h4 style="margin-top:0;color:#003087;">How to Use Film &amp; Grade</h4>
        <ol style="color:#1F2937;line-height:2.0;">
        <li><b>Upload photos</b> of the play — screenshots, stills, or a phone photo.
            Video clips work too, but photos are faster and more reliable.</li>
        <li><b>Load Media</b> — everything becomes analyzable images</li>
        <li><b>Pick a mode</b> — ask a specific question, or run a scored RefGrade</li>
        <li><b>Review</b> — every answer cites NFHS rules and starts with a
            Visibility Check</li>
        <li><b>Export</b> to TXT or PDF</li>
        </ol>
        <p style="font-size:0.82rem;color:#4B5563;margin-bottom:0;">
        <em>Tip: three or four well-chosen stills (pre-snap, at the snap, at the point
        of the foul) usually beat a full video clip for both speed and cost.</em>
        </p>
        </div>""", unsafe_allow_html=True)



# ─────────────────────────────────────────────────────────────────────────────
# TAB 3 — REF HUB  (v1.1: Pre-Game first, Crew + Ref Eval merged)
#
# Two sections instead of three. Pre-Game Meeting leads because it is the most
# frequently used and needs no upload. Crew Eval and Ref Eval were near-identical
# forms differing only in scope, so they are now one Evaluations section with a
# scope dropdown — "Full Crew" routes to CREW_EVAL_PROMPT, any single position
# routes to REF_EVAL_PROMPT. Both accept photos as well as video.
# ─────────────────────────────────────────────────────────────────────────────

with tab_ah:
    st.markdown("## 👥 Ref Hub")
    st.markdown(
        "Auto-generated pre-game meeting agendas, plus film- and photo-based "
        "crew and individual official evaluations."
    )

    # ── Section selector ─────────────────────────────────────────────────────
    sub_a, sub_b = st.columns(2)
    with sub_a:
        _active = st.session_state.ah_sub == "pregame"
        if st.button("📅 Pre-Game Meeting", use_container_width=True,
                     key="ah_sub_pregame",
                     type="primary" if _active else "secondary"):
            st.session_state.ah_sub = "pregame"
            st.rerun()
    with sub_b:
        _active = st.session_state.ah_sub == "eval"
        if st.button("🎬 Evaluations", use_container_width=True,
                     key="ah_sub_eval",
                     type="primary" if _active else "secondary"):
            st.session_state.ah_sub = "eval"
            st.rerun()

    st.markdown("---")

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 1 — PRE-GAME MEETING
    # ═══════════════════════════════════════════════════════════════════════
    if st.session_state.ah_sub == "pregame":
        st.markdown("### 📅 Pre-Game Meeting Agenda Generator")
        st.markdown(
            "Auto-generates a pre-game agenda from the RefBuddy Knowledge Base, "
            "including 2026 rule changes, key mechanics, and your own additional "
            "notes section for a fully customized meeting."
        )

        pg1, pg2 = st.columns(2)
        with pg1:
            pg_crew = st.selectbox("Crew size",
                                    ["3-Person Crew", "4-Person Crew", "5-Person Crew"],
                                    key="pg_crew_sel")
            pg_level = st.selectbox("Game level",
                                     ["Varsity", "Junior Varsity", "9th Grade", "Playoff"],
                                     key="pg_level_sel")
        with pg2:
            pg_date = st.text_input("Game date (optional)",
                                     placeholder="e.g. Friday, September 4",
                                     key="pg_date")
            pg_teams = st.text_input("Teams (optional)",
                                      placeholder="e.g. Eden Prairie vs Wayzata",
                                      key="pg_teams")

        pg_focus = st.multiselect(
            "Additional emphasis topics (optional)",
            options=["2026 Rule Changes", "Mercy Rule Procedure", "Overtime Procedure",
                     "Targeting/Defenseless Players", "Onside Kick Readiness",
                     "Goal Line Mechanics", "Penalty Reporting", "Equipment Checks",
                     "12-Player Situations", "Play Clock Management",
                     "Fair Catch Interference", "Forward Fumble Rule Change"],
            key="pg_focus_sel",
        )

        pg_assignor_notes = st.text_area(
            "Assignor's Custom Notes / Emphasis",
            height=130,
            placeholder="Add anything you want to emphasize for this game",
            key="pg_assignor_notes",
        )

        if st.button("📅  Generate Pre-Game Meeting Agenda",
                     type="primary", use_container_width=True, key="pg_generate"):
            focus_str = (f"Additional emphasis topics requested: {', '.join(pg_focus)}\n"
                         if pg_focus else "")
            header_str = ""
            if pg_date or pg_teams:
                header_str = (f"Game: {pg_teams or 'TBD'} | "
                              f"Date: {pg_date or 'TBD'} | {pg_level}\n")

            # Preserve line breaks from the notes box as separate bullets so
            # multi-line notes are not run together in the agenda
            if pg_assignor_notes.strip():
                raw_lines = [l.strip() for l in pg_assignor_notes.strip().splitlines()
                             if l.strip()]
                if len(raw_lines) == 1:
                    notes_section = raw_lines[0]
                else:
                    notes_section = "\n".join(f"- {ln.lstrip('-* ').strip()}"
                                               for ln in raw_lines)
            else:
                notes_section = "(No specific assignor notes provided for this game.)"

            prompt = (
                f"Generate a pre-game meeting agenda for the following game.\n\n"
                f"{header_str}"
                f"Crew configuration: {pg_crew}\n"
                f"Game level: {pg_level}\n"
                f"{focus_str}\n"
                f"For the Assignor Notes section, use EXACTLY this content verbatim — "
                f"do not summarize or rephrase it:\n{notes_section}\n\n"
                f"Generate the full agenda following your system prompt structure."
            )

            with st.spinner("Generating pre-game meeting agenda… (15–30 seconds)"):
                try:
                    result = call_api_sync(prompt, PREGAME_MEETING_PROMPT, max_tokens=3000)
                    st.session_state.ah_pregame_result = result
                    st.session_state.ah_pregame_logs.append({
                        "timestamp": datetime.datetime.now().isoformat(),
                        "teams": pg_teams or "Unknown", "date": pg_date or "Unknown",
                        "crew": pg_crew, "level": pg_level, "result": result,
                    })
                except Exception as e:
                    st.error(handle_api_error(e))

        if st.session_state.ah_pregame_result:
            st.markdown("---")
            st.markdown("### Generated Agenda")
            with st.expander("📋 View Full Agenda", expanded=True):
                st.markdown(st.session_state.ah_pregame_result)

            st.markdown("**Export**")
            x1, x2, x3, x4 = st.columns(4)
            ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            safe = (st.session_state.get("pg_teams") or "game").replace(" ", "_")[:25]
            title = f"Pre-Game Meeting Agenda — {st.session_state.get('pg_teams') or 'Game'}"
            with x1:
                st.download_button("⬇️ TXT", data=st.session_state.ah_pregame_result,
                                   file_name=f"pregame_{safe}_{ts}.txt",
                                   mime="text/plain", use_container_width=True)
            with x2:
                _p = markdown_to_pdf_bytes(st.session_state.ah_pregame_result, title)
                if _p:
                    st.download_button("⬇️ PDF", data=_p,
                                       file_name=f"pregame_{safe}_{ts}.pdf",
                                       mime="application/pdf", use_container_width=True)
            with x3:
                _d = markdown_to_docx_bytes(st.session_state.ah_pregame_result, title)
                if _d:
                    st.download_button(
                        "⬇️ Word", data=_d, file_name=f"pregame_{safe}_{ts}.docx",
                        mime=("application/vnd.openxmlformats-officedocument"
                              ".wordprocessingml.document"),
                        use_container_width=True)
            with x4:
                if st.button("🗑️ Clear", use_container_width=True, key="pg_clear"):
                    st.session_state.ah_pregame_result = ""
                    st.rerun()

    # ═══════════════════════════════════════════════════════════════════════
    # SECTION 2 — EVALUATIONS  (crew + individual, photos or video)
    # ═══════════════════════════════════════════════════════════════════════
    elif st.session_state.ah_sub == "eval":
        st.markdown("### 🎬 Crew & Official Evaluations")
        st.markdown(
            "Upload photos or a clip, choose whether you're evaluating the whole "
            "crew or one official, and RefBuddy produces a scored report with "
            "rule citations and coaching points."
        )

        ev_uploads = st.file_uploader(
            "ev_upload",
            type=["jpg", "jpeg", "png", "mp4", "mov"],
            accept_multiple_files=True,
            label_visibility="collapsed",
            key="ah_ev_uploader",
        )
        st.caption("📸 Photos (JPG/PNG) or 🎥 video (MP4/MOV). Photos upload faster "
                   "and are usually enough — try pre-snap, at the snap, and the "
                   "moment of the foul.")

        e1, e2 = st.columns(2)
        with e1:
            ev_scope = st.selectbox(
                "Evaluate",
                options=["Full Crew (Overall)", "Referee (R)", "Umpire (U)",
                         "Line Judge (LJ)", "Down Judge (DJ)", "Back Judge (BJ)",
                         "Side Judge (SJ)", "Field Judge (FJ)"],
                key="ev_scope",
            )
        with e2:
            ev_crew = st.selectbox(
                "Crew configuration",
                ["3-Person Crew", "4-Person Crew", "5-Person Crew"],
                key="ev_crew",
            )

        ev_fps = 1.0
        if ev_uploads and any(f.name.lower().endswith((".mp4", ".mov")) for f in ev_uploads):
            ev_fps = st.select_slider(
                "Video sampling rate (frames per second)",
                options=[0.5, 1.0, 2.0], value=1.0, key="ev_fps",
            )

        ev_notes = st.text_area(
            "Notes / focus (optional)", height=90,
            placeholder="e.g. 'Check BJ depth on the punt.' or "
                        "'A flag was thrown here — was it correct?'",
            key="ev_notes",
        )

        is_crew = ev_scope.startswith("Full Crew")
        btn_label = ("🎬  Generate Crew Evaluation" if is_crew
                     else f"🏈  Generate {ev_scope} Evaluation")

        if st.button(btn_label, type="primary", use_container_width=True,
                     disabled=not ev_uploads, key="ev_run"):
            frames = []
            with st.spinner("Processing upload…"):
                for f in ev_uploads:
                    if f.name.lower().endswith((".jpg", ".jpeg", ".png")):
                        b = image_to_frame_b64(f)
                        if b:
                            frames.append(b)
                    else:
                        if not OPENCV_AVAILABLE:
                            st.error("Video requires opencv-python-headless.")
                            break
                        try:
                            sfx = ".mp4" if f.name.lower().endswith(".mp4") else ".mov"
                            with tempfile.NamedTemporaryFile(delete=False, suffix=sfx) as t:
                                t.write(f.read()); tp = t.name
                            frames.extend(extract_frames(tp, fps=ev_fps))
                            os.unlink(tp)
                        except Exception as e:
                            st.error(f"❌ Could not process {f.name}: {e}")

            if not frames:
                st.error("Nothing could be read from that upload.")
            else:
                cap = min(len(frames), 40)
                if spend_frames(cap):
                    name = ", ".join(f.name for f in ev_uploads[:3])
                    extra = f"\nNotes: {ev_notes.strip()}" if ev_notes.strip() else ""
                    if is_crew:
                        q = (f"Perform a full crew evaluation.\nSource: {name}\n"
                             f"Crew configuration: {ev_crew}\n"
                             f"Images analyzed: 1–{cap} of {len(frames)}{extra}\n\n"
                             f"Analyze all visible officials for positioning, call "
                             f"accuracy, mechanics, dead-ball officiating and "
                             f"communication. Begin with a VISIBILITY CHECK.")
                        prompt_used = CREW_EVAL_PROMPT
                        heading = "📊 Crew Evaluation Report"
                    else:
                        q = (f"Evaluate ONLY the {ev_scope} in this material.\n"
                             f"Source: {name}\nCrew configuration: {ev_crew}\n"
                             f"Images analyzed: 1–{cap} of {len(frames)}{extra}\n\n"
                             f"Focus entirely on this one official. Ignore others unless "
                             f"their actions directly affect this official's "
                             f"responsibilities. Begin with a VISIBILITY CHECK for this "
                             f"position only.")
                        prompt_used = REF_EVAL_PROMPT
                        heading = f"📊 {ev_scope} Evaluation Report"

                    blocks = build_vision_content(
                        frames, 0, cap - 1, q, name, ev_fps,
                        preamble_extra=("These may be still photos rather than "
                                        "sequential video frames. Visibility Check "
                                        "is the mandatory first section."),
                    )
                    st.markdown("---")
                    st.markdown(f"#### {heading}")
                    ph = st.empty(); full = ""
                    try:
                        with st.spinner(f"Analyzing {cap} image(s)… (30–120 seconds)"):
                            for chunk in stream_vision(make_client(), blocks, prompt_used):
                                full += chunk; ph.markdown(full + "▌")
                        ph.markdown(full)
                        st.session_state.ah_eval_result = full
                        st.session_state.ah_eval_scope = ev_scope
                    except Exception as e:
                        st.error(handle_api_error(e))

        if st.session_state.get("ah_eval_result"):
            st.markdown("---")
            with st.expander("📄 Evaluation Report", expanded=True):
                st.markdown(st.session_state.ah_eval_result)

            st.markdown("**Export**")
            z1, z2, z3 = st.columns(3)
            ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            with z1:
                st.download_button("⬇️ TXT", data=st.session_state.ah_eval_result,
                                   file_name=f"evaluation_{ts}.txt",
                                   mime="text/plain", use_container_width=True)
            with z2:
                _p = markdown_to_pdf_bytes(
                    st.session_state.ah_eval_result,
                    f"{st.session_state.get('ah_eval_scope','')} Evaluation Report")
                if _p:
                    st.download_button("⬇️ PDF", data=_p,
                                       file_name=f"evaluation_{ts}.pdf",
                                       mime="application/pdf",
                                       use_container_width=True)
            with z3:
                if st.button("🗑️ Clear", use_container_width=True, key="ev_clear"):
                    st.session_state.ah_eval_result = ""
                    st.rerun()



# ─────────────────────────────────────────────────────────────────────────────

with tab_quiz:
    st.markdown("## 📝 Quiz")
    st.markdown("Test your football officiating knowledge! Questions are generated "
                "from the RefBuddy Knowledge Base with a 50/50 mix of multiple-choice "
                "and true/false.")


    # ── MODE SELECTOR ─────────────────────────────────────────────────────────
    st.markdown("### Choose Your Mode")
    mc1, mc2 = st.columns(2)

    with mc1:
        is_practice = st.session_state.quiz_mode == "practice"
        card_cls = "mode-card-active" if is_practice else "mode-card-inactive"
        st.markdown(f"""<div class="{card_cls}">
        <div style="font-size:1.25rem;font-weight:800;color:{BLUE};">🔁 One-by-One Practice</div>
        <div style="color:{MUTED};font-size:0.88rem;margin-top:0.3rem;">
        Unlimited questions · instant feedback · live accuracy · no repeats</div>
        </div>""", unsafe_allow_html=True)
        if st.button("Select Practice Mode", key="sel_practice", use_container_width=True):
            st.session_state.quiz_mode = "practice"
            st.session_state.quiz_current_q = None
            st.session_state.quiz_answered = False
            st.session_state.quiz_user_answer = None
            st.session_state.quiz_total = 0
            st.session_state.quiz_correct = 0
            st.session_state.quiz_session_topics = []
            st.rerun()

    with mc2:
        is_tenq = st.session_state.quiz_mode == "tenq"
        card_cls2 = "mode-card-active" if is_tenq else "mode-card-inactive"
        st.markdown(f"""<div class="{card_cls2}">
        <div style="font-size:1.25rem;font-weight:800;color:{BLUE};">🎯 10-Question Quiz</div>
        <div style="color:{MUTED};font-size:0.88rem;margin-top:0.3rem;">
        Set of 10 · progress bar · final score · full review + save</div>
        </div>""", unsafe_allow_html=True)
        if st.button("Select 10-Question Quiz", key="sel_tenq", use_container_width=True):
            st.session_state.quiz_mode = "tenq"
            st.session_state.tenq_questions = []
            st.session_state.tenq_index = 0
            st.session_state.tenq_answers = []
            st.session_state.tenq_finished = False
            st.session_state.tenq_answered_this = False
            st.session_state.tenq_user_answer = None
            st.rerun()

    # ── TOPIC ─────────────────────────────────────────────────────────────────
    if st.session_state.quiz_mode:
        st.markdown("---")
        st.markdown("**Topic Focus**")
        topics = ["Mixed", "Rules", "Mechanics", "Positioning",
                  "Signals", "Game Situations", "MSHSL Specific", "2026 Changes"]
        new_topic = st.selectbox("topic_sel", topics,
                                  index=topics.index(st.session_state.quiz_topic),
                                  label_visibility="collapsed", key="topic_selector")
        if new_topic != st.session_state.quiz_topic:
            st.session_state.quiz_topic = new_topic
            st.session_state.quiz_current_q = None
            st.session_state.quiz_answered = False
            st.session_state.quiz_session_topics = []
            st.session_state.tenq_questions = []
            st.session_state.tenq_index = 0
            st.session_state.tenq_answers = []
            st.session_state.tenq_finished = False
            st.rerun()

    st.markdown("---")

    # ═══════════════════════════════════════════════════════════════════════════
    # PRACTICE MODE
    # ═══════════════════════════════════════════════════════════════════════════

    if st.session_state.quiz_mode == "practice":
        st.markdown("### 🔁 One-by-One Practice")

        if st.session_state.quiz_total > 0:
            accuracy_display(st.session_state.quiz_correct, st.session_state.quiz_total)

        # Generate first question if none loaded
        if st.session_state.quiz_current_q is None:
            with st.spinner("Generating your question…"):
                # Try up to 3 times (in case of wrong option count)
                q = None
                for _ in range(3):
                    q = generate_single_question(
                        st.session_state.quiz_topic,
                        st.session_state.quiz_session_topics
                    )
                    if q is not None:
                        break
            if q:
                st.session_state.quiz_current_q = q
                st.session_state.quiz_answered = False
                st.session_state.quiz_user_answer = None
                # Track topic to avoid repeats
                topic_key = q.get("topic", "") + ": " + q.get("question", "")[:50]
                st.session_state.quiz_session_topics.append(topic_key)
                st.rerun()
            else:
                st.error("Failed to generate a question. Check your API key and try again.")
                st.stop()

        q = st.session_state.quiz_current_q

        render_question_card(q, question_num=f"**Q{st.session_state.quiz_total + 1}.**")

        if not st.session_state.quiz_answered:
            options = q.get("options", {})
            option_labels = [f"{k}:  {v}" for k, v in sorted(options.items())]
            user_choice = st.radio(
                "**Select your answer:**",
                options=option_labels,
                key=f"prac_radio_{st.session_state.quiz_total}",
            )
            if st.button("✅ Submit Answer", use_container_width=True, key="prac_submit"):
                chosen = user_choice.split(":")[0].strip()
                st.session_state.quiz_user_answer = chosen
                st.session_state.quiz_answered = True
                st.session_state.quiz_total += 1
                if render_feedback(q, chosen):
                    st.session_state.quiz_correct += 1
                st.rerun()

        else:
            render_feedback(q, st.session_state.quiz_user_answer)
            topic_tag = q.get("topic", "")
            if topic_tag:
                st.markdown(f'<span class="pill-blue">📁 {topic_tag}</span>', unsafe_allow_html=True)
            st.markdown("")

            nav1, nav2 = st.columns([3, 1])
            with nav1:
                if st.button("➡️ Next Question", use_container_width=True, key="prac_next"):
                    with st.spinner("Generating next question…"):
                        q_new = None
                        for _ in range(3):
                            q_new = generate_single_question(
                                st.session_state.quiz_topic,
                                st.session_state.quiz_session_topics
                            )
                            if q_new is not None:
                                break
                    if q_new:
                        st.session_state.quiz_current_q = q_new
                        st.session_state.quiz_answered = False
                        st.session_state.quiz_user_answer = None
                        topic_key = q_new.get("topic", "") + ": " + q_new.get("question", "")[:50]
                        st.session_state.quiz_session_topics.append(topic_key)
                        if len(st.session_state.quiz_session_topics) > 20:
                            st.session_state.quiz_session_topics = \
                                st.session_state.quiz_session_topics[-10:]
                        st.rerun()
                    else:
                        st.error("Failed to generate next question.")
            with nav2:
                if st.button("🔄 Reset", use_container_width=True, key="prac_reset"):
                    st.session_state.quiz_current_q = None
                    st.session_state.quiz_answered = False
                    st.session_state.quiz_user_answer = None
                    st.session_state.quiz_total = 0
                    st.session_state.quiz_correct = 0
                    st.session_state.quiz_session_topics = []
                    st.rerun()

    # ═══════════════════════════════════════════════════════════════════════════
    # 10-QUESTION QUIZ MODE
    # ═══════════════════════════════════════════════════════════════════════════

    elif st.session_state.quiz_mode == "tenq":
        st.markdown("### 🎯 10-Question Quiz")

        if not st.session_state.tenq_questions and not st.session_state.tenq_finished:
            st.markdown("""<div class="rb-card" style="text-align:center;padding:1.5rem;">
            <p style="color:#1F2937;margin-bottom:1rem;">
            Click <strong>Generate My Quiz</strong> to get 10 questions (5 MC + 5 True/False)
            covering your selected topic.
            </p></div>""", unsafe_allow_html=True)

            if st.button("🚀 Generate My 10-Question Quiz",
                         use_container_width=True, key="tenq_gen"):
                with st.spinner("Generating 10 questions (5 MC + 5 True/False)… 15–30 seconds"):
                    qs = generate_ten_questions(st.session_state.quiz_topic)
                if qs:
                    st.session_state.tenq_questions = qs
                    st.session_state.tenq_index = 0
                    st.session_state.tenq_answers = []
                    st.session_state.tenq_finished = False
                    st.session_state.tenq_answered_this = False
                    st.session_state.tenq_user_answer = None
                    st.rerun()
                else:
                    st.error("Failed to generate quiz. Check your API key and try again.")

        elif st.session_state.tenq_questions and not st.session_state.tenq_finished:
            idx = st.session_state.tenq_index
            total_qs = len(st.session_state.tenq_questions)
            q = st.session_state.tenq_questions[idx]

            st.progress(idx / total_qs)
            st.markdown(f'<div style="text-align:right;font-size:0.85rem;color:{MUTED};">'
                        f'Question {idx+1} of {total_qs}</div>', unsafe_allow_html=True)

            render_question_card(q, question_num=f"**Q{idx+1}.**")
            options = q.get("options", {})

            if not st.session_state.tenq_answered_this:
                option_labels = [f"{k}:  {v}" for k, v in sorted(options.items())]
                user_choice = st.radio("**Select your answer:**", options=option_labels,
                                       key=f"tenq_radio_{idx}")
                if st.button("✅ Submit Answer", use_container_width=True,
                             key=f"tenq_submit_{idx}"):
                    chosen = user_choice.split(":")[0].strip()
                    st.session_state.tenq_user_answer = chosen
                    st.session_state.tenq_answered_this = True
                    is_correct = chosen == q.get("correct", "")
                    st.session_state.tenq_answers.append({
                        "question_num": idx + 1,
                        "user": chosen, "correct": q.get("correct", ""),
                        "is_correct": is_correct, "data": q,
                    })
                    st.rerun()
            else:
                render_feedback(q, st.session_state.tenq_user_answer)
                st.markdown("")
                is_last = (idx == total_qs - 1)
                btn_lbl = "📊 See Final Score" if is_last else f"➡️ Next ({idx+2}/{total_qs})"
                if st.button(btn_lbl, use_container_width=True, key=f"tenq_next_{idx}"):
                    if is_last:
                        st.session_state.tenq_finished = True
                    else:
                        st.session_state.tenq_index += 1
                        st.session_state.tenq_answered_this = False
                        st.session_state.tenq_user_answer = None
                    st.rerun()

        elif st.session_state.tenq_finished and st.session_state.tenq_answers:
            answers = st.session_state.tenq_answers
            n_correct = sum(1 for a in answers if a["is_correct"])
            n_total = len(answers)
            pct = int(round(n_correct / n_total * 100))
            score_color = ("#15803D" if pct >= 80 else ("#92400E" if pct >= 60 else "#991B1B"))
            grade_label = ("🏆 Excellent!" if pct >= 90 else "✅ Good" if pct >= 80
                           else "📈 Getting there" if pct >= 70 else "📚 Keep studying"
                           if pct >= 60 else "🔁 Review the material")

            st.markdown(f"""
            <div style="background:{CARD};border:2px solid {score_color};border-radius:14px;
                        padding:2rem;text-align:center;margin-bottom:1.5rem;
                        box-shadow:0 4px 16px rgba(0,0,0,0.08);">
                <div style="font-size:3.5rem;font-weight:900;color:{score_color};">{pct}%</div>
                <div style="font-size:1.3rem;font-weight:700;color:#1F2937;margin:0.3rem 0;">
                    {n_correct} / {n_total} correct &nbsp; {grade_label}</div>
                <div style="color:{MUTED};font-size:0.9rem;">Topic: {st.session_state.quiz_topic}</div>
            </div>""", unsafe_allow_html=True)

            ra1, ra2 = st.columns(2)
            with ra1:
                if st.button("📁 Save Results to My Log", use_container_width=True, key="tenq_save"):
                    st.session_state.quiz_log.append({
                        "timestamp": datetime.datetime.now().isoformat(),
                        "topic": st.session_state.quiz_topic,
                        "score": pct, "correct": n_correct, "total": n_total,
                        "answers": answers,
                    })
                    st.success(f"✅ Saved! {len(st.session_state.quiz_log)} quiz log(s) on file.")
            with ra2:
                if st.button("🔄 Take Another Quiz", use_container_width=True, key="tenq_restart"):
                    st.session_state.tenq_questions = []
                    st.session_state.tenq_index = 0
                    st.session_state.tenq_answers = []
                    st.session_state.tenq_finished = False
                    st.session_state.tenq_answered_this = False
                    st.session_state.tenq_user_answer = None
                    st.rerun()

            st.markdown("---")
            st.markdown("### 📋 Full Review")
            for a in answers:
                qd = a["data"]
                opts = qd.get("options", {})
                u, c, ic = a["user"], a["correct"], a["is_correct"]
                icon = "✅" if ic else "❌"
                cbg = "#F0FDF4" if ic else "#FFF1F2"
                cbo = "#4ADE80" if ic else "#F87171"
                u_txt, c_txt = opts.get(u, u), opts.get(c, c)
                corr_line = (
                    "" if ic
                    else f'<br><strong style="color:#7F1D1D;">✔ Correct: {c}: {c_txt}</strong>'
                )
                st.markdown(f"""
                <div style="background:{cbg};border:1.5px solid {cbo};border-radius:10px;
                            padding:1.1rem 1.3rem;margin-bottom:0.9rem;">
                    <div style="font-weight:700;color:#1F2937;">
                        {icon} Q{a["question_num"]}: {qd.get("question","")}</div>
                    <div style="font-size:0.9rem;color:#1F2937;margin-top:0.3rem;">
                        <strong>Your answer:</strong> {u}: {u_txt}{corr_line}</div>
                </div>""", unsafe_allow_html=True)
                with st.expander(f"📖 Explanation — Q{a['question_num']}", expanded=False):
                    p = qd.get("personal_note", "")
                    pnote = f'<br><strong>📋 From your notes:</strong> {p}' if p else ""
                    st.markdown(f"""<div class="quiz-explanation">
                    {qd.get("explanation","")}<br><br>
                    <strong>📌 Citation:</strong> {qd.get("rule_citation","")}{pnote}
                    </div>""", unsafe_allow_html=True)

            if st.session_state.quiz_log:
                st.markdown("---")
                ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                st.download_button(
                    f"⬇️ Download All Quiz Results ({len(st.session_state.quiz_log)} saved)",
                    data=json.dumps(st.session_state.quiz_log, indent=2, ensure_ascii=False),
                    file_name=f"refbuddy_quiz_{ts}.json",
                    mime="application/json", use_container_width=True,
                )

    elif st.session_state.quiz_mode is None:
        st.markdown("""<div class="rb-card" style="text-align:center;padding:1.5rem;">
        <p style="color:#4B5563;margin:0;">👆 Select a mode above to get started.</p>
        </div>""", unsafe_allow_html=True)


# =============================================================================
# FOOTER
# =============================================================================

st.markdown(f"""
<div class="rb-footer">
    Built by a ref, for refs 🏈 &nbsp;|&nbsp;
    Years of NFHS veteran officials&rsquo; game notes and NFHS/MSHSL rulebook facts and interpretations<br>
    <span style="font-size:0.72rem;">
    Always confirm rulings with your MSHSL assignor. Not official NFHS/MSHSL interpretation.
    </span>
</div>
""", unsafe_allow_html=True)

# ── Terms of Use / Disclaimer (v3.2) ─────────────────────────────────────────
with st.expander("📄 Terms of Use, Disclaimer & Attribution", expanded=False):
    st.markdown("""
#### Not an official source
RefBuddy is an independent study and preparation tool built by a working MSHSL
football official. **It is not affiliated with, endorsed by, sponsored by, or
approved by the National Federation of State High School Associations (NFHS),
the Minnesota State High School League (MSHSL), or any officials' association.**
NFHS and MSHSL names are used here only to identify the rule sets being discussed.

#### No rulings, no substitute for the rulebook
Nothing in RefBuddy is an official rule interpretation. Answers are generated by
an AI model and **can be wrong, incomplete, or out of date.** RefBuddy summarizes
rules in plain language and points you to rule numbers — it does not reproduce
rulebook text. **You are expected to own and consult a current NFHS Football
Rules Book and the current MSHSL Minnesota modifications.** For any ruling that
matters, confirm with your district assignor or state rules interpreter before
acting on it.

#### Intellectual property
Rule numbers, penalty distances, timing values, and mechanics positions are
factual information. The NFHS Rules Book, Case Book, Points of Emphasis, and
MSHSL publications are copyrighted works owned by their respective organizations
and are **not** reproduced here. Personal game notes and all original commentary
are the author's own work.

#### Your uploads
Video clips and files you upload are held in memory for the length of your
session, sent to the Anthropic API for analysis, and are not stored on a server
by this app. Do not upload anything containing sensitive personal information
about students, players, coaches, or officials. Chat logs and reports you
download are saved only to your own device.

#### Acceptable use
Access is limited to officials who have been given the crew password. Don't
share it publicly. Usage is capped per session to keep costs sustainable.

#### No warranty
This tool is provided "as is," without warranty of any kind. The author is not
liable for any outcome resulting from its use, including missed or incorrect
calls, evaluations, or game administration decisions.
    """)
